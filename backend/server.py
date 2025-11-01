def add_subscription_endpoint(router):
    @router.get("/auth/subscription")
    async def get_subscription(request: Request):
        """Return dummy subscription status (inactive) for all users."""
        auth_header = request.headers.get("Authorization")
        if not auth_header or not auth_header.startswith("Bearer "):
            return {"active": False, "plan": None, "expires": None}
    
        return {"active": False, "plan": None, "expires": None}

from fastapi import FastAPI, APIRouter, File, UploadFile, HTTPException, Form, Depends, Request, Body
from fastapi.responses import FileResponse, StreamingResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
try:
    from gtts import gTTS
except Exception:
    gTTS = None
import io
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
from pymongo import MongoClient
import os
import logging
from datetime import datetime, timedelta
from bson import ObjectId

from fastapi import APIRouter, HTTPException
import pymongo
import os
import json
import logging
from pathlib import Path
from datetime import datetime, timezone
from typing import List, Dict, Any


# Set up paths
ROOT_DIR = Path(__file__).parent

# Configure logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

# Load environment variables first
load_dotenv(ROOT_DIR / '.env')
logging.info("Environment variables loaded")

# Log critical configuration
mongo_url = os.environ.get('MONGODB_URL') or os.environ.get('MONGO_URL')
if mongo_url:
    # Log URL without credentials
    safe_url = mongo_url.split('@')[-1] if '@' in mongo_url else 'mongodb://****'
    logging.info(f"MONGODB_URL found: mongodb://****@{safe_url}")
else:
    logging.warning("No MONGODB_URL or MONGO_URL found in environment")

from pathlib import Path
from pydantic import BaseModel, Field
from typing import List, Optional, Any
from dataclasses import dataclass
import uuid
from datetime import datetime, timezone, timedelta
import shutil
import httpx
import asyncio
from PyPDF2 import PdfReader
try:
    # pdfminer for more robust PDF text extraction
    from pdfminer.high_level import extract_text as pdfminer_extract_text
except Exception:
    pdfminer_extract_text = None
try:
    # python-docx for .docx extraction
    import docx
except Exception:
    docx = None
import hashlib
from passlib.context import CryptContext
import bcrypt
from jwt import encode as jwt_encode, decode as jwt_decode  # Import from PyJWT
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
import re
import secrets
import requests
import unicodedata
import tempfile
import json
import difflib
from io import BytesIO

import os
import re
import uuid
import logging
import pymongo
from fastapi import APIRouter, HTTPException
from PyPDF2 import PdfReader
from datetime import datetime, timezone





def strip_markdown(text: str) -> str:
    """Remove common Markdown formatting to produce clean plain text."""
    try:
        if not text:
            return text
        # Remove code fences ``` ```
        text = re.sub(r'```[\s\S]*?```', '', text)
        # Replace inline code `code` with code
        text = re.sub(r'`([^`]*)`', r'\1', text)
        # Remove ATX headings (# Headline)
        text = re.sub(r'^\s{0,3}#{1,6}\s*', '', text, flags=re.MULTILINE)
        # Replace bold/italic markers while preserving inner text
        text = re.sub(r'\*\*(.*?)\*\*', r"\1", text)
        text = re.sub(r'__(.*?)__', r"\1", text)
        text = re.sub(r'\*(.*?)\*', r"\1", text)
        text = re.sub(r'_(.*?)_', r"\1", text)
        # Remove list markers at start of lines
        text = re.sub(r'^[\s]*[-\*\+]\s+', '', text, flags=re.MULTILINE)
        # Remove remaining stray asterisks
        text = text.replace('*', '')
        # Collapse multiple blank lines
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()
    except Exception:
        return text


def send_email_via_resend(to_email: str, subject: str, html_body: str, text_body: str, from_email: str = None) -> bool:
    """Send an email using Resend (https://resend.com/) REST API.

    Requires environment variable RESEND_API_KEY. Returns True on success, False otherwise.
    """
    try:
        api_key = os.environ.get('RESEND_API_KEY')
        if not api_key:
            logging.debug("RESEND_API_KEY not set, cannot send via Resend")
            return False

        # Always use no-reply@waterbears.in as sender
        sender = 'no-reply@waterbears.in'
        payload = {
            "from": sender,
            "to": [to_email],
            "subject": subject,
            "html": html_body,
            "text": text_body
        }

        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }

        resp = requests.post("https://api.resend.com/emails", json=payload, headers=headers, timeout=15)
        if resp.status_code in (200, 202):
            logging.info(f"Resend: email sent to {to_email}, status={resp.status_code}")
            return True
        else:
            logging.error(f"Resend email failed: status={resp.status_code} body={resp.text}")
            return False
    except Exception as e:
        logging.error(f"Resend email exception: {e}")
        return False

# Authentication imports
from passlib.context import CryptContext
from datetime import timedelta
from fastapi.security import HTTPBearer

# Google OAuth configuration
GOOGLE_CLIENT_ID = os.environ.get('GOOGLE_CLIENT_ID')
GOOGLE_CLIENT_SECRET = os.environ.get('GOOGLE_CLIENT_SECRET')
GOOGLE_REDIRECT_URI = os.environ.get('GOOGLE_REDIRECT_URI', 'http://localhost:3000/auth/google/callback')

# Custom classes for message handling
@dataclass
class UserMessage:
    text: str
    file_contents: Optional[List[Any]] = None

@dataclass
class FileContentWithMimeType:
    file_path: str
    mime_type: str

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

logging.basicConfig(level=logging.INFO)
logging.info(f"GEMINI_API_KEY loaded: {'set' if os.environ.get('GEMINI_API_KEY') else 'NOT set'}")
if not os.environ.get('GEMINI_API_KEY'):
    logging.error("GEMINI_API_KEY is not set. Document analysis will not work!")

# Create the main app without a prefix
app = FastAPI()

# Serve static files and assets from frontend build
build_path = ROOT_DIR.parent / "frontend" / "build"
static_path = build_path / "static"
if build_path.exists():
    # Serve the static directory
    app.mount("/static", StaticFiles(directory=str(static_path)), name="static")
    # Serve other files from the build directory
    app.mount("/", StaticFiles(directory=str(build_path), html=True), name="frontend")

@app.exception_handler(404)
async def not_found_handler(request: Request, exc: HTTPException):
    if not str(request.url.path).startswith("/api/"):
        index_html = build_path / "index.html"
        if index_html.exists():
            return FileResponse(str(index_html))
    return JSONResponse(status_code=404, content={"detail": "Not found"})

# MongoDB helper function
def get_mongo_client():
    """Get a MongoDB client with connection verification"""
    mongo_url = os.environ.get('MONGODB_URL') or os.environ.get('MONGO_URL', 'mongodb://localhost:27017')
    try:
        client = pymongo.MongoClient(mongo_url, serverSelectionTimeoutMS=5000)
        # Force a connection to verify it works
        client.server_info()
        # Log successful connection without exposing credentials
        safe_url = mongo_url.split('@')[0] + '@****' if '@' in mongo_url else 'mongodb://****'
        logging.info(f"MongoDB connection successful to {safe_url}")
        return client
    except pymongo.errors.ServerSelectionTimeoutError as e:
        logging.error(f"MongoDB connection timeout: {str(e)}")
        raise HTTPException(status_code=500, detail="Database connection failed - timeout")
    except pymongo.errors.ConnectionFailure as e:
        logging.error(f"MongoDB connection failed: {str(e)}")
        raise HTTPException(status_code=500, detail="Database connection failed")
    except Exception as e:
        logging.error(f"MongoDB error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Database error: {str(e)}")

# MongoDB connection state
class MongoDBState:
    async_client = None
    async_db = None
    sync_client = None
    sync_db = None

    @classmethod
    async def initialize(cls):
        try:
            # Get MongoDB URL from environment variables - Railway provides MONGODB_URL
            mongo_url = os.environ.get('MONGODB_URL')
            if not mongo_url:
                mongo_url = os.environ.get('MONGO_URL')
                
            if not mongo_url:
                logging.error("No MongoDB connection URL found in environment variables")
                logging.error("Checked both MONGODB_URL and MONGO_URL")
                logging.error("Available environment variables: " + ", ".join(sorted(os.environ.keys())))
                raise Exception("MongoDB connection URL not found in environment")
                
            # Log URL without exposing credentials
            safe_url = mongo_url.split('@')[-1] if '@' in mongo_url else 'mongodb://****'
            logging.info(f"Initializing MongoDB connection to: mongodb://****@{safe_url}")
            
            # Configure MongoDB clients with optimized settings
            client_settings = {
                'serverSelectionTimeoutMS': 5000,
                'heartbeatFrequencyMS': 30000,
                'retryWrites': True,
                'connectTimeoutMS': 30000,
                'socketTimeoutMS': 30000,
                'maxIdleTimeMS': 60000,
                'maxPoolSize': 10,
                'retryReads': True,
                'w': 'majority',
                'journal': True,
                'appname': 'CovenantAI'
            }
            
            # Initialize async client
            cls.async_client = AsyncIOMotorClient(mongo_url, **client_settings)
            
            # Initialize sync client with same settings
            cls.sync_client = MongoClient(mongo_url, **client_settings)
            
            db_name = os.environ.get('DB_NAME', 'legal_docs')
            cls.async_db = cls.async_client[db_name]
            cls.sync_db = cls.sync_client[db_name]
            
            # Test connection
            await cls.async_client.admin.command('ping')
            cls.sync_client.admin.command('ping')
            logging.info(f"MongoDB connected successfully to database: {db_name}")
            
            # Create indexes
            await cls.create_indexes()
            
        except Exception as e:
            logging.error(f"MongoDB connection failed: {str(e)}")
            cls.async_client = None
            cls.async_db = None
            cls.sync_client = None
            cls.sync_db = None
            raise

    @classmethod
    async def create_indexes(cls):
        if cls.async_db is not None:
            try:
                await cls.async_db.documents.create_index([("user_id", 1)])
                await cls.async_db.documents.create_index([("upload_date", -1)])
                await cls.async_db.chat_messages.create_index([("user_id", 1)])
                await cls.async_db.chat_messages.create_index([("timestamp", -1)])
                await cls.async_db.users.create_index([("email", 1)], unique=True)
                logging.info("MongoDB indexes created successfully")
            except Exception as e:
                logging.error(f"Failed to create MongoDB indexes: {str(e)}")

    @classmethod
    async def check_connection(cls):
        while True:
            try:
                if cls.async_client is None or cls.async_db is None or cls.sync_client is None or cls.sync_db is None:
                    logging.warning("Attempting to reconnect to MongoDB...")
                    await cls.initialize()
                else:
                    # Test both connections
                    await cls.async_client.admin.command('ping')
                    cls.sync_client.admin.command('ping')
                await asyncio.sleep(30)
            except Exception as e:
                logging.error(f"MongoDB connection check failed: {str(e)}")
                await asyncio.sleep(5)

    @classmethod
    def cleanup(cls):
        """Clean up MongoDB connections on shutdown"""
        if cls.async_client:
            cls.async_client.close()
        if cls.sync_client:
            cls.sync_client.close()

# Initialize MongoDB on startup
@app.on_event("startup")
async def startup_mongodb():
    try:
        await MongoDBState.initialize()
        # Start connection monitoring in background
        asyncio.create_task(MongoDBState.check_connection())
    except Exception as e:
        logging.error(f"Failed to initialize MongoDB: {str(e)}")

# Add cleanup on shutdown
@app.on_event("shutdown")
async def shutdown_mongodb():
    MongoDBState.cleanup()

# Database access helpers
def get_db(async_db=True):
    """Get database handle. Returns async or sync db based on parameter."""
    if async_db:
        if MongoDBState.async_db is None:
            raise HTTPException(status_code=503, detail="Database connection not available")
        return MongoDBState.async_db
    else:
        if MongoDBState.sync_db is None:
            raise HTTPException(status_code=503, detail="Database connection not available")
        return MongoDBState.sync_db

# Periodic connection check
async def check_mongodb_connection():
    while True:
        try:
            if (
                MongoDBState.async_client is None or
                MongoDBState.async_db is None or
                MongoDBState.sync_client is None or
                MongoDBState.sync_db is None
            ):
                logging.warning("Attempting to reconnect to MongoDB...")
                await MongoDBState.initialize()
            else:
                await MongoDBState.async_client.admin.command('ping')
                MongoDBState.sync_client.admin.command('ping')
            await asyncio.sleep(30)  # Check every 30 seconds
        except Exception as e:
            logging.error(f"MongoDB connection check failed: {str(e)}")
            await asyncio.sleep(5)  # Wait before retry

# Start connection monitoring
@app.on_event("startup")
async def start_mongodb_monitor():
    asyncio.create_task(check_mongodb_connection())

# CORS Configuration
# In production, use CORS_ORIGINS environment variable to specify allowed origins
# In development, allow localhost origins
cors_origins_env = os.environ.get('CORS_ORIGINS')
railway_url = os.environ.get('RAILWAY_STATIC_URL')  # Railway provides this

origins = [
    "https://covenant.up.railway.app",
    "https://backendcovenantai.up.railway.app"
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")
add_subscription_endpoint(api_router)

# Create uploads directory
UPLOAD_DIR = ROOT_DIR / "../uploads"
UPLOAD_DIR.mkdir(exist_ok=True)

# Authentication setup with enhanced security and error handling
pwd_context = CryptContext(
    schemes=["bcrypt"],
    deprecated="auto",
    bcrypt__default_rounds=12,
    bcrypt__truncate_error=False
)

# Get secret key from environment or generate a secure one
SECRET_KEY = os.environ.get('SECRET_KEY')
if not SECRET_KEY:
    logging.warning("SECRET_KEY not set in environment - generating a secure random key")
    SECRET_KEY = secrets.token_urlsafe(32)
    logging.info("Generated SECRET_KEY (save this for production): " + SECRET_KEY)

# JWT Configuration
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = int(os.environ.get('ACCESS_TOKEN_EXPIRE_MINUTES', '30'))
MIN_PASSWORD_LENGTH = 8  # Minimum password length requirement

def hash_password(password: str) -> str:
    """Hash a password for storing using bcrypt directly.

    bcrypt enforces a 72-byte limit on the input. We truncate the UTF-8
    encoded bytes to 72 bytes, then hash using bcrypt.gensalt(). The resulting
    hash is ASCII-safe, so we store it as a utf-8 decoded string.
    """
    pw_bytes = password.encode('utf-8')
    if len(pw_bytes) > 72:
        logging.warning("Password exceeds 72 bytes, it will be truncated to 72 bytes before hashing")
        pw_bytes = pw_bytes[:72]

    hashed = bcrypt.hashpw(pw_bytes, bcrypt.gensalt(rounds=12))
    return hashed.decode('utf-8')

def create_access_token(data: dict, expires_delta: timedelta = None):
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=15)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt_encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

def verify_token(token: str):
    """Verify a JWT access token and return the decoded payload as a dict.

    On success: returns the payload (dict).
    On failure: raises HTTPException(status_code=401) so calling routes receive an auth error.
    """
    if not token:
        raise HTTPException(status_code=401, detail="Authorization token missing")

    try:
        payload = jwt_decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        if not isinstance(payload, dict):
            raise HTTPException(status_code=401, detail="Invalid token payload")
        return payload
    except Exception as e:
        if "expired" in str(e).lower():
            logging.debug("JWT expired")
            raise HTTPException(status_code=401, detail="Token expired")
        logging.debug(f"JWT decode error: {e}")
        raise HTTPException(status_code=401, detail="Invalid authentication token")

# Pydantic Models
class Document(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    filename: str
    file_path: str
    file_type: str
    upload_date: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    analysis_status: str = "pending"  # pending, processing, completed, failed
    content: Optional[str] = None  # Extracted document text
    metadata: Optional[dict] = None  # Storage for content stats and extraction info
    summary: Optional[str] = None  # Executive summary
    key_points: Optional[List[dict]] = None  # Key points extracted from the document
    key_clauses: Optional[List[dict]] = None  # Important clauses found
    risk_assessment: Optional[dict] = None  # Structured risk assessment
    plain_english: Optional[str] = None  # Plain English explanation
    recommendations: Optional[List[dict]] = None  # Actionable recommendations
    analysis_version: str = "2.0"  # Track analysis system version

class DocumentCreate(BaseModel):
    filename: str
    file_type: str

class QuestionRequest(BaseModel):
    document_id: str
    question: str
    session_id: Optional[str] = None

class ChatMessage(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    document_id: str
    session_id: str
    question: str
    answer: str
    timestamp: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class GlobalChatMessage(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    session_id: str
    user_id: Optional[str] = None  # None for anonymous users
    chat_mode: str = "general"  # "general" or "document"
    selected_document: Optional[str] = None
    question: str
    answer: str
    document_metadata: Optional[dict] = None  # Store relevant document context
    timestamp: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class AnalysisRequest(BaseModel):
    document_id: str

class ExportRequest(BaseModel):
    sections: List[str]

# User models for authentication
class User(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    email: str
    name: str
    google_id: Optional[str] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class LoginRequest(BaseModel):
    email: str
    password: str

class RegisterRequest(BaseModel):
    email: str
    password: str
    name: str

class TokenResponse(BaseModel):
    access_token: str
    token_type: str = "bearer"
    user: dict

# Document text extraction system
def validate_pdf(file_path: str) -> bool:
    """Validate PDF file integrity"""
    try:
        with open(file_path, 'rb') as file:
            # Check PDF header
            header = file.read(5)
            if header != b'%PDF-':
                return False
            
            # Try to load with PyPDF2 to check structure
            reader = PdfReader(file_path)
            if len(reader.pages) == 0:
                return False
                
            return True
    except Exception as e:
        logging.warning(f"PDF validation failed: {e}")
        return False

def extract_document_text(file_path: str, file_type: str) -> tuple[str, str]:
    """
    Enhanced document text extraction with multiple fallback methods.
    Returns: (extracted_text, method_used)
    """
    logging.info(f"Starting document extraction for {file_path} of type {file_type}")
    
    text = ""
    method = ""
    error_messages = []
    
    # For PDFs, validate first
    if file_type == 'application/pdf' or file_path.lower().endswith('.pdf'):
        if not validate_pdf(file_path):
            error_messages.append("Invalid or corrupted PDF file")
            return "Invalid or corrupted PDF file", "failed"
    
    # PDF Handling
    if file_type == 'application/pdf' or file_path.lower().endswith('.pdf'):
        # 1. Try pdfminer.six first (more robust)
        try:
            from pdfminer.high_level import extract_text
            from pdfminer.layout import LAParams
            
            # Configure LAParams for better text extraction
            laparams = LAParams(
                line_margin=0.5,
                word_margin=0.1,
                char_margin=2.0,
                boxes_flow=0.5,
                detect_vertical=True
            )
            
            pdfminer_text = extract_text(file_path, laparams=laparams)
            if pdfminer_text and pdfminer_text.strip():
                text = pdfminer_text
                # Clean up extracted text
                text = re.sub(r'\s+', ' ', text)  # Normalize whitespace
                text = re.sub(r'\n\s*\n', '\n\n', text)  # Fix paragraph breaks
                text = text.strip()
                
                if len(text) > 100:  # Check if we got meaningful text
                    method = "pdfminer"
                    logging.info(f"pdfminer extraction successful: {len(text)} chars")
                    return text, method
        except Exception as e:
            error_messages.append(f"pdfminer failed: {str(e)}")
            
        # 2. Try PyPDF2 as backup
        try:
            reader = PdfReader(file_path)
            pages = []
            for i, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        pages.append(page_text)
                        logging.info(f"PyPDF2 extracted page {i}: {len(page_text)} chars")
                except Exception as e:
                    logging.warning(f"PyPDF2 failed on page {i}: {e}")
            
            pypdf_text = "\n".join(pages)
            if pypdf_text.strip() and len(pypdf_text) > 100:
                text = pypdf_text
                # Clean up extracted text
                text = re.sub(r'\s+', ' ', text)
                text = re.sub(r'\n\s*\n', '\n\n', text)
                text = text.strip()
                method = "PyPDF2"
                logging.info(f"PyPDF2 extraction successful: {len(text)} chars")
                return text, method
        except Exception as e:
            error_messages.append(f"PyPDF2 failed: {str(e)}")
            
        # 2. Try pdfminer.six
        try:
            if pdfminer_extract_text:
                pdfminer_text = pdfminer_extract_text(file_path)
                if pdfminer_text and len(pdfminer_text.strip()) > len(text.strip()):
                    text = pdfminer_text
                    method = "pdfminer"
                    logging.info(f"pdfminer extraction successful: {len(text)} chars")
                    return text, method
        except Exception as e:
            error_messages.append(f"pdfminer failed: {str(e)}")
            
        # 3. Try OCR as last resort
        if len(text.strip()) < 100:
            try:
                import pytesseract
                from PIL import Image
                from pdf2image import convert_from_path
                import os
                
                logging.info("Attempting OCR extraction")
                ocr_text = []
                
                # Configure tesseract for better accuracy
                custom_config = r'--oem 3 --psm 6'
                
                with tempfile.TemporaryDirectory() as tempdir:
                    try:
                        # Convert PDF to images with higher DPI for better OCR
                        images = convert_from_path(file_path, output_folder=tempdir, dpi=300)
                        
                        for i, image in enumerate(images):
                            # Preprocess image for better OCR
                            # Convert to RGB if not already
                            if image.mode != 'RGB':
                                image = image.convert('RGB')
                            
                            # Try OCR with custom config
                            page_text = pytesseract.image_to_string(image, config=custom_config)
                            
                            if page_text.strip():
                                # Clean up OCR text
                                cleaned_text = re.sub(r'\s+', ' ', page_text)  # Normalize whitespace
                                cleaned_text = re.sub(r'\n\s*\n', '\n\n', cleaned_text)  # Fix paragraph breaks
                                ocr_text.append(cleaned_text.strip())
                                logging.info(f"OCR extracted page {i}: {len(cleaned_text)} chars")
                    except Exception as e:
                        logging.error(f"Error during PDF to image conversion: {e}")
                        
                final_ocr_text = "\n\n".join(ocr_text)
                if final_ocr_text and len(final_ocr_text.strip()) > len(text.strip()):
                    text = final_ocr_text
                    method = "OCR"
                    logging.info(f"OCR extraction successful: {len(text)} chars")
                    return text, method
            except Exception as e:
                error_messages.append(f"OCR failed: {str(e)}")
                
    # DOCX Handling
    elif file_path.lower().endswith('.docx'):
        try:
            if docx:
                doc = docx.Document(file_path)
                paragraphs = []
                for p in doc.paragraphs:
                    if p.text.strip():
                        paragraphs.append(p.text)
                text = "\n".join(paragraphs)
                if text.strip():
                    method = "python-docx"
                    logging.info(f"DOCX extraction successful: {len(text)} chars")
                    return text, method
        except Exception as e:
            error_messages.append(f"DOCX extraction failed: {str(e)}")
    
    # Text file handling
    else:
        # Try UTF-8 first
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
                if text.strip():
                    method = "utf-8"
                    logging.info(f"UTF-8 read successful: {len(text)} chars")
                    return text, method
        except UnicodeDecodeError:
            # Try latin-1 as fallback
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    text = f.read()
                    if text.strip():
                        method = "latin-1"
                        logging.info(f"Latin-1 read successful: {len(text)} chars")
                        return text, method
            except Exception as e:
                error_messages.append(f"Text read failed: {str(e)}")
    
    # If we reach here, all methods failed or text is empty
    if not text.strip():
        error_detail = "; ".join(error_messages) if error_messages else "No valid text content found"
        logging.error(f"All extraction methods failed: {error_detail}")
        return f"Document text extraction failed: {error_detail}", "failed"
    
    return text, method

def extract_text_pdf_fallback(file_path: str) -> str:
    """Attempt to extract PDF text using pdfminer.six as a fallback."""
    try:
        if pdfminer_extract_text:
            logging.info(f"Attempting pdfminer.six extraction on {file_path}")
            try:
                text = pdfminer_extract_text(file_path)
                logging.info(f"pdfminer.six extraction result length: {len(text) if text else 0}")
                if text:
                    logging.debug(f"pdfminer.six sample: {text[:200]}")
                return text or ""
            except Exception as e:
                logging.warning(f"pdfminer extraction failed: {e}")
        else:
            logging.info("pdfminer.six not available in environment")
    except Exception as e:
        logging.error(f"Unexpected pdfminer fallback error: {e}")
    return ""

def extract_text_docx(file_path: str) -> str:
    """Extract text from .docx files using python-docx."""
    try:
        if docx:
            logging.info(f"Attempting python-docx extraction for .docx: {file_path}")
            try:
                document = docx.Document(file_path)
                paragraphs = [p.text for p in document.paragraphs if p.text and p.text.strip()]
                logging.info(f"python-docx extraction result length: {sum(len(p) for p in paragraphs)}")
                if paragraphs:
                    logging.debug(f"python-docx sample: {paragraphs[0][:200]}")
                return "\n".join(paragraphs)
            except Exception as e:
                logging.warning(f"python-docx extraction failed: {e}")
        else:
            logging.info("python-docx not available in environment")
    except Exception as e:
        logging.error(f"Unexpected docx fallback error: {e}")
    return ""

GEMINI_API_KEY = os.getenv('GEMINI_API_KEY', 'AIzaSyCuBibZH55rvDR8b8Utty9ThK_hUTBh3Es')
GEMINI_API_URL = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={GEMINI_API_KEY}"

async def send_message(prompt: str, mode: str = "general", context: str = "") -> str:
    logging.debug(f"Starting AI request with prompt length: {len(prompt)}, mode: {mode}")
    
    # Set up the system message based on mode
    if mode == "document":
        system_prefix = """You are a legal document analysis assistant. Your task is to:
1. Answer questions based ONLY on the provided document content
2. If information isn't in the document, clearly state that
3. Quote relevant sections when possible
4. Keep responses concise and focused
5. Do not make assumptions beyond the document content

"""
        full_prompt = system_prefix + prompt
    else:
        full_prompt = prompt

    async def local_fallback(p: str) -> str:
        """Generate a safe, concise fallback response when the external AI is unavailable.

        This provides reasonable behavior for development and offline testing so the
        `/api/chat` and document analysis endpoints remain functional.
        """
        # If the prompt looks like a long document, return a short structured analysis
        text = p.strip()
        if len(text) > 2000 or '\n' in text and len(text) > 800:
            # Very short heuristic analysis: executive  + key points
            summary = text[:800].strip()
            return (
                "EXECUTIVE : This document appears to be a legal text. "
                "A quick read indicates the main topics and obligations are: "
                f"{summary[:800]}...\n\nKEY POINTS: (1) Check parties and dates; (2) Look for termination and liability clauses; (3) Identify obligations and deliverables."
            )

        # If it's a question-style prompt, reply concisely
        if text.endswith('?') or text.lower().startswith(('what', 'how', 'why', 'explain', 'when', 'where', 'who')):
            return "I'm a local fallback assistant: I can give quick answers, but the full AI service isn't available. Try a concise question or enable the AI API for detailed responses."

        # Default echo-style friendly reply
        short = text[:600]
        return f"Acknowledged. Here's a brief response based on your input: {short}{'...' if len(text) > 400 else ''}"

    # Try external Gemini API if configured; otherwise use local fallback
    if GEMINI_API_KEY:
        max_retries = 4
        for attempt in range(max_retries):
            try:
                async with httpx.AsyncClient(timeout=30.0) as client:
                    response = await client.post(
                        GEMINI_API_URL,
                        json={
                            "contents": [{"parts": [{"text": prompt}]}]
                        },
                        headers={"Content-Type": "application/json"},
                    )
                    if response.status_code == 200:
                        data = response.json()
                        logging.debug(f"Gemini API response keys: {list(data.keys()) if isinstance(data, dict) else type(data)}")
                        try:
                            candidate = data.get('candidates', [])[0]
                            response_text = candidate.get('content', {}).get('parts', [])[0].get('text', '')
                            if response_text:
                                return response_text
                        except Exception:
                            logging.debug('Unexpected Gemini response structure, falling back')
                    else:
                        logging.warning(f"Gemini status {response.status_code}; body: {response.text}")
            except Exception as e:
                logging.warning(f"Gemini API attempt {attempt+1} failed: {e}")

    # External API disabled or failed -> return local fallback
    logging.info('Using local AI fallback response')
    # local_fallback is async; await it directly instead of passing the coroutine into
    # run_in_executor (which expects a synchronous callable). Awaiting guarantees a
    # concrete string is returned and avoids "coroutine was never awaited" warnings.
    return await local_fallback(prompt)

@api_router.get("/")
async def root():
    return {"message": "Legal Document AI Assistant API"}

@api_router.post("/auth/login")
async def login(request: LoginRequest):
    """Login or auto-register user and send verification if needed"""
    try:
        sync_client = None
        # logging is already imported globally

        logging.info(f"Login attempt for email: {request.email}")

        db = get_db(async_db=False)  # Use synchronous DB
        user = db.users.find_one({"email": request.email})
        logging.info(f"Found user: {user is not None}")
        if user:
            logging.info(f"User ID: {user.get('id', user.get('_id', 'MISSING'))}")
            logging.info(f"Email verified: {user.get('email_verified')}")
            logging.info(f"User document: {user}")
        else:
            logging.warning(f"No user found for email: {request.email}")
        stored_password = user.get('password', '') if user else ''

        # Validate user fields
        if user:
            if 'password' not in user or not user['password']:
                logging.error(f"User record missing password field: {user}")
                raise HTTPException(status_code=500, detail="User record missing password.")
            if 'email_verified' not in user:
                logging.error(f"User record missing email_verified field: {user}")
                raise HTTPException(status_code=500, detail="User record missing email_verified.")
            if 'email_verify_code' not in user:
                logging.warning(f"User record missing email_verify_code field: {user}")
            if 'email_verify_expires' not in user:
                logging.warning(f"User record missing email_verify_expires field: {user}")

        # 🧩 If user does not exist → auto-register and send verification code
        if not user:
            code = secrets.token_hex(3)
            expires = (datetime.utcnow() + timedelta(minutes=15)).isoformat()

            hashed_pw = bcrypt.hashpw(request.password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

            new_user = {
                "email": request.email,
                "password": hashed_pw,
                "email_verified": False,
                "email_verify_code": code,
                "email_verify_expires": expires,
                "created_at": datetime.utcnow().isoformat(),
            }

            inserted = db.users.insert_one(new_user)
            user_id = str(inserted.inserted_id)

            # --- send verification email ---
            smtp_host = os.environ.get('SMTP_HOST')
            smtp_port = int(os.environ.get('SMTP_PORT', '587'))
            smtp_user = os.environ.get('SMTP_USER')
            smtp_pass = os.environ.get('SMTP_PASS')
            from_email = os.environ.get('SMTP_FROM', 'CovenantAI <no-reply@waterbears.in>')

            subject = "Verify Your Email - CovenantAI"
            text_body = f"""Hi there,

Please verify your email using this code: {code}

This code will expire in 15 minutes.

Best,
CovenantAI Team
"""

            html_body = f"""
            <div style="font-family: 'Segoe UI', sans-serif; background-color: #f4f4f7; padding: 20px;">
                <div style="max-width: 600px; margin: auto; background: #fff; border-radius: 10px;
                            box-shadow: 0 3px 8px rgba(0,0,0,0.05); overflow: hidden;">
                    <div style="background: linear-gradient(135deg, #3b82f6, #2563eb); color: white; padding: 18px; text-align: center;">
                        <h2 style="margin: 0;">Verify Your Email</h2>
                    </div>
                    <div style="padding: 25px; color: #333;">
                        <p>Hi <strong>{request.email}</strong>,</p>
                        <p>Thanks for joining <b>CovenantAI</b>! Please verify your email using the code below:</p>
                        <div style="text-align: center; margin: 30px 0;">
                            <div style="display: inline-block; background: #2563eb; color: white; 
                                        font-size: 20px; letter-spacing: 3px; padding: 12px 24px; 
                                        border-radius: 6px;">
                                {code}
                            </div>
                        </div>
                        <p>This code will expire in <b>15 minutes</b>.</p>
                        <p>If you didn’t create an account, ignore this email.</p>
                        <hr style="border:none; border-top:1px solid #eee; margin: 25px 0;">
                        <p style="font-size: 13px; color: #777;">© {datetime.utcnow().year} CovenantAI. All rights reserved.</p>
                    </div>
                </div>
            </div>
            """

            # Send via Resend only. If Resend is not configured, fall back to logging/printing the code.
            if os.environ.get('RESEND_API_KEY'):
                try:
                    sent = send_email_via_resend(request.email, subject, html_body, text_body, from_email=from_email)
                    if sent:
                        logging.info(f"Verification email sent via Resend to {request.email}")
                    else:
                        logging.error(f"Resend failed to send verification to {request.email}")
                        logging.info(f"Verification code for {request.email}: {code}")
                except Exception as e:
                    logging.error(f"Error sending via Resend: {e}")
                    logging.info(f"Verification code for {request.email}: {code}")
            else:
                logging.info(f"Resend not configured; verification code for {request.email}: {code}")

            return {
                "message": "Verification code sent. Please verify your email before login.",
                "email": request.email,
                "requires_verification": True
            }

        # Existing user → verify password
        try:
            logging.info(f"Verifying password for user: {request.email}")
            logging.info(f"Stored password: {stored_password[:20] if stored_password else 'EMPTY'}...")
            logging.info(f"Input password: {request.password[:5]}...")
            
            # Add debug logging
            logging.info("Password verification details:")
            logging.info(f"Input password length: {len(request.password)}")
            logging.info(f"Stored hash length: {len(stored_password)}")
            logging.info(f"Stored hash type: {type(stored_password)}")
            logging.info(f"Full user record: {user}")
            logging.info(f"Checking bcrypt password...")

            # Prepare bytes for bcrypt
            password_bytes = request.password.encode('utf-8')
            hash_bytes = stored_password.encode('utf-8') if isinstance(stored_password, str) else stored_password
            
            logging.info(f"Password bytes: {password_bytes[:20]}...")
            logging.info(f"Hash bytes: {hash_bytes[:20]}...")

            valid = bcrypt.checkpw(password_bytes, hash_bytes)
            logging.info(f"Password valid: {valid}")

            if not valid:
                logging.error(f"Password check failed for {request.email}")
                raise HTTPException(status_code=401, detail="Invalid credentials")

        except HTTPException:
            # Re-raise HTTP exceptions
            raise
        except Exception as e:
            logging.error(f"Password verification failed: {str(e)}", exc_info=True)
            raise HTTPException(status_code=401, detail="Invalid credentials")

        # If user email not verified → resend code if needed
        if not user.get('email_verified', False):
            code = user.get('email_verify_code') or secrets.token_hex(3)
            expires = (datetime.utcnow() + timedelta(minutes=15)).isoformat()

            db.users.update_one(
                {"_id": user["_id"]},
                {"$set": {
                    "email_verify_code": code,
                    "email_verify_expires": expires
                }}
            )

            # Re-send verification email
            smtp_host = os.environ.get('SMTP_HOST')
            smtp_port = int(os.environ.get('SMTP_PORT', '587'))
            smtp_user = os.environ.get('SMTP_USER')
            smtp_pass = os.environ.get('SMTP_PASS')
            from_email = os.environ.get('SMTP_FROM', 'CovenantAI <no-reply@waterbears.in>')

            subject = "Verify Your Email - CovenantAI"
            text_body = f"""Hi {user['email']},

Please verify your email using this code: {code}

This code will expire in 15 minutes.

Best,
CovenantAI Team
"""

            html_body = f"""
            <div style="font-family: 'Segoe UI', sans-serif; background-color: #f4f4f7; padding: 20px;">
                <div style="max-width: 600px; margin: auto; background: #fff; border-radius: 10px;
                            box-shadow: 0 3px 8px rgba(0,0,0,0.05); overflow: hidden;">
                    <div style="background: linear-gradient(135deg, #3b82f6, #2563eb); color: white; padding: 18px; text-align: center;">
                        <h2 style="margin: 0;">Verify Your Email</h2>
                    </div>
                    <div style="padding: 25px; color: #333;">
                        <p>Hi <strong>{user['email']}</strong>,</p>
                        <p>Use the code below to verify your account:</p>
                        <div style="text-align: center; margin: 30px 0;">
                            <div style="display: inline-block; background: #2563eb; color: white; 
                                        font-size: 20px; letter-spacing: 3px; padding: 12px 24px; 
                                        border-radius: 6px;">
                                {code}
                            </div>
                        </div>
                        <p>This code will expire in <b>15 minutes</b>.</p>
                        <p>If you didn’t request this, ignore this email.</p>
                        <hr style="border:none; border-top:1px solid #eee; margin: 25px 0;">
                        <p style="font-size: 13px; color: #777;">© {datetime.utcnow().year} CovenantAI. All rights reserved.</p>
                    </div>
                </div>
            </div>
            """

            # Send via Resend only. If Resend is not configured, log the verification code.
            if os.environ.get('RESEND_API_KEY'):
                try:
                    sent = send_email_via_resend(user['email'], "Verify Your Email - CovenantAI", html_body, text_body, from_email=from_email)
                    if sent:
                        logging.info(f"Verification email re-sent via Resend to {user['email']}")
                    else:
                        logging.error(f"Resend failed to re-send verification to {user['email']}")
                        logging.info(f"Verification code for {user['email']}: {code}")
                except Exception as e:
                    logging.error(f"Resend error when re-sending verification: {e}")
                    logging.info(f"Verification code for {user['email']}: {code}")
            else:
                logging.info(f"Resend not configured; verification code for {user['email']}: {code}")

            return {
                "message": "Email verification required",
                "email": user['email'],
                "requires_verification": True
            }

        # If verified → create access token and login
        user_id = user.get('id') or str(user.get('_id'))  # Fallback to string _id if no id field
        logging.info(f"Creating access token for user ID: {user_id}")
        
        token_data = {
            "sub": user['email'], 
            "user_id": user_id,
            "email": user['email']
        }
        access_token = create_access_token(
            data=token_data,
            expires_delta=timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
        )
        logging.info(f"Access token created successfully")

    # No sync_client to close here

        # Create standardized user response
        user_response = {
            "id": user_id,
            "email": user['email'],
            "name": user.get('name', '')
        }
        logging.info(f"Returning successful login response")
        
        return TokenResponse(
            access_token=access_token,
            user=user_response
        )

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Login error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Login failed")

@api_router.get("/history")
async def get_history(request: Request):
    """Get user's document history"""
    try:
        # Extract token from Authorization header
        auth_header = request.headers.get("Authorization")
        logging.info(f"/history called - Authorization header present: {bool(auth_header)}")
        if not auth_header or not auth_header.startswith("Bearer "):
            logging.warning("/history: missing or invalid Authorization header")
            raise HTTPException(status_code=401, detail="Authorization header missing or invalid")

        token = auth_header.replace("Bearer ", "")
        payload = verify_token(token)
        logging.info(f"/history: token verification payload present: {bool(payload)}")
        if not payload:
            logging.warning("/history: token verification failed or payload empty")
            raise HTTPException(status_code=401, detail="Invalid token")

        user_id = payload.get("user_id")
        logging.info(f"/history: extracted user_id: {user_id}")
        if not user_id:
            logging.warning("/history: token payload missing user_id")
            raise HTTPException(status_code=401, detail="Invalid token payload")

        # Get user's documents
        mongo_url = os.environ.get('MONGO_URL', 'mongodb://localhost:27017')
        logging.info(f"/history: connecting to MongoDB at: {('****@' + mongo_url.split('@')[-1]) if '@' in mongo_url else mongo_url}")
        sync_client = pymongo.MongoClient(mongo_url)
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
        query = {"user_id": user_id}
        logging.info(f"/history: query={query}")
        documents = list(sync_db.documents.find(query).sort("created_at", -1))
        logging.info(f"/history: documents found={len(documents)}")
        if len(documents) > 0:
            try:
                sample_ids = [str(d.get('id') or d.get('_id')) for d in documents[:5]]
                logging.info(f"/history: sample ids: {sample_ids}")
            except Exception:
                logging.debug("/history: could not extract sample ids from documents")
        sync_client.close()

        # Format documents for response, robust to missing/legacy fields
        history = []
        for doc in documents:
            doc_id = doc.get('id') or str(doc.get('_id', ''))
            analysis_result = doc.get('analysis_result', {})
            if isinstance(analysis_result, str):
                try:
                    analysis_result = json.loads(analysis_result)
                except Exception:
                    analysis_result = {}
            created_at = doc.get('created_at', '')
            if isinstance(created_at, datetime):
                created_at = created_at.isoformat()
            else:
                created_at = str(created_at)
            history.append({
                'id': doc_id,
                'filename': doc.get('filename', 'Untitled Document'),
                'document_text': doc.get('document_text', '')[:200] + '...' if len(doc.get('document_text', '')) > 200 else doc.get('document_text', ''),
                'analysis_result': analysis_result,
                'summary': doc.get('summary', ''),
                'risk_score': doc.get('risk_score', {}),
                'critical_flags': doc.get('critical_flags', []),
                'analysis_status': doc.get('analysis_status', 'completed'),
                'created_at': created_at
            })

        return history

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Get history error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to retrieve history")




def safe_json(obj):
    """Recursively make MongoDB objects (ObjectId, datetime, etc.) JSON serializable"""
    if isinstance(obj, ObjectId):
        return str(obj)
    elif isinstance(obj, datetime):
        return obj.isoformat()
    elif isinstance(obj, dict):
        return {k: safe_json(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [safe_json(v) for v in obj]
    return obj


@api_router.get("/history/guest")
async def get_guest_history():
    """Get guest document history (documents with no user_id)"""
    try:
        # --- 1. Connect to MongoDB ---
        mongo_url = os.environ.get("MONGO_URL", "mongodb://localhost:27017")
        db_name = os.environ.get("DB_NAME", "legal_docs")
        logging.info(f"/history/guest: connecting to MongoDB at: {('****@' + mongo_url.split('@')[-1]) if '@' in mongo_url else mongo_url}")
        sync_client = pymongo.MongoClient(mongo_url)
        sync_db = sync_client[db_name]

        # --- 2. Fetch guest (unauthenticated) documents ---
        documents = list(sync_db.documents.find({
            "$or": [
                {"user_id": {"$exists": False}},
                {"user_id": None},
                {"user_id": ""}
            ]
        }).sort("upload_date", -1))  # Sort newest first

        # --- 3. Close client ---
        logging.info(f"/history/guest: documents count={len(documents)}")
        if len(documents) > 0:
            try:
                sample = [str(d.get('id') or d.get('_id')) for d in documents[:5]]
                logging.info(f"/history/guest: sample ids: {sample}")
            except Exception:
                logging.debug("/history/guest: could not extract sample ids")
        sync_client.close()

        # --- 4. Format for response ---
        history = []
        for doc in documents:
            doc_id = doc.get('id') or str(doc.get('_id', ''))
            analysis_result = doc.get('analysis_result', {})
            if isinstance(analysis_result, str):
                try:
                    analysis_result = json.loads(analysis_result)
                except Exception:
                    analysis_result = {}
            upload_date_val = doc.get("upload_date")
            upload_dt = None
            if isinstance(upload_date_val, datetime):
                upload_dt = upload_date_val
            elif isinstance(upload_date_val, str):
                try:
                    upload_dt = datetime.fromisoformat(upload_date_val)
                except Exception:
                    upload_dt = None
            expires_at = (upload_dt + timedelta(minutes=10)).isoformat() if upload_dt else None
            upload_time_str = upload_dt.isoformat() if upload_dt else (
                str(upload_date_val) if upload_date_val is not None else ""
            )
            history.append({
                "id": doc_id,
                "filename": doc.get("filename", "Untitled Document"),
                "analysis_result": safe_json(analysis_result),
                "upload_time": upload_time_str,
                "expires_at": expires_at
            })
        return history

    except Exception as e:
        import traceback
        logging.error(f"Get guest history error: {e}\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Failed to retrieve guest history: {str(e)}")



@api_router.post("/history")
async def save_to_history(request: Request, data: dict = Body(...)):
    """Save document to history"""
    try:
        logging.info(f"🔵 /history POST endpoint called")
        logging.info(f"   Request data keys: {list(data.keys())}")
        # Extract token from Authorization header
        auth_header = request.headers.get("Authorization")
        logging.info(f"   Authorization header present: {bool(auth_header)}")

        if not auth_header or not auth_header.startswith("Bearer "):
            logging.error("Authorization header missing or invalid")
            raise HTTPException(status_code=401, detail="Authorization header missing or invalid")

        token = auth_header.replace("Bearer ", "")
        payload = verify_token(token)
        logging.info(f"   token verification payload present: {bool(payload)}")
        if not payload:
            logging.error("Invalid token")
            raise HTTPException(status_code=401, detail="Invalid token")

        user_id = payload.get("user_id")
        if not user_id:
            logging.error("Invalid token payload - no user_id")
            raise HTTPException(status_code=401, detail="Invalid token payload")

        logging.info(f"   User ID: {user_id}")

        document_text = data.get('documentText')
        analysis_result = data.get('analysisResult', {})
        filename = data.get('filename', 'Untitled Document')

        logging.info(f"   Document text length: {len(document_text) if document_text else 0}")
        logging.info(f"   Analysis result keys: {list(analysis_result.keys()) if analysis_result else []}")
        logging.info(f"   Filename: {filename}")

        if not document_text:
            logging.error("Document text is required")
            raise HTTPException(status_code=400, detail="Document text is required")

        if not analysis_result:
            logging.error("Analysis result is required")
            raise HTTPException(status_code=400, detail="Analysis result is required")

        # Save to database
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]

        doc_data = {
            "id": str(uuid.uuid4()),
            "user_id": user_id,
            "filename": filename,
            "document_text": document_text,
            "analysis_result": analysis_result,
            "created_at": datetime.now(timezone.utc),
            "summary": analysis_result.get('summary', ''),
            "risk_score": analysis_result.get('riskScore', {}),
            "critical_flags": analysis_result.get('criticalFlags', []),
            "analysis_status": "completed"
        }

        logging.info(f"   Inserting document: {doc_data.get('id')}")
        sync_db.documents.insert_one(doc_data)
        sync_client.close()
        
        logging.info(f"✅ Document saved to history successfully: {doc_data.get('id')}")

        return {"message": "Document saved to history"}

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Save to history error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to save to history")

@api_router.post("/chat")
async def chat_endpoint(request: Request, data: dict = Body(...)):
    """Chat with AI about legal questions - works for authenticated and anonymous users"""
    logging.info(f"🤖 CHAT ENDPOINT CALLED - Data: {data}")
    sync_client = None
    try:
        # Extract token from Authorization header (optional for anonymous users)
        user_id = None
        auth_header = request.headers.get("Authorization")
        logging.info(f"🤖 Auth header: {auth_header if auth_header else 'NONE'}")
        if auth_header and auth_header.startswith("Bearer "):
            logging.info(f"🤖 Bearer token found, extracting...")
            token = auth_header.replace("Bearer ", "")
            logging.info(f"🤖 Token: {token[:50]}...")
            try:
                payload = verify_token(token)
                logging.info(f"🤖 verify_token returned: {payload}")
                if payload:
                    user_id = payload.get("user_id")
                    logging.info(f"🤖 User ID extracted: {user_id}")
                else:
                    logging.info(f"🤖 verify_token returned None/False")
            except Exception as e:
                logging.error(f"🤖 verify_token EXCEPTION: {e}", exc_info=True)

        message = data.get('message')
        context = data.get('context', [])
        mode = data.get('mode', 'general')  # Get chat mode from request
        selected_document = data.get('selected_document')

        if not message:
            raise HTTPException(status_code=400, detail="Message is required")

        # Create MongoDB client for document retrieval
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]

        # Build context from previous messages and document if in document mode
        document_content_for_ai = ""
        if mode == "document" and selected_document:
            try:
                # Get document content from database
                logging.info(f"Fetching document {selected_document} for chat")
                document = sync_db.documents.find_one({"id": selected_document})
                if not document:
                    raise HTTPException(status_code=404, detail="Document not found")

                # Try different content fields in order of preference
                document_content_for_ai = (
                    document.get('content', '') or 
                    document.get('document_text', '') or
                    document.get('summary', '')
                )

                # Fallback: If no content, try to read from file
                if not document_content_for_ai:
                    file_path = document.get('file_path')
                    file_type = document.get('file_type', '')
                    if file_path:
                        from pathlib import Path
                        # Try the file_path as stored
                        if Path(file_path).exists():
                            try:
                                if file_type == 'application/pdf' or str(file_path).lower().endswith('.pdf'):
                                    # Extract from PDF
                                    from PyPDF2 import PdfReader
                                    reader = PdfReader(file_path)
                                    pages = []
                                    for page in reader.pages:
                                        text = page.extract_text()
                                        if text:
                                            pages.append(text)
                                    document_content_for_ai = "\n".join(pages)
                                    logging.info(f"Extracted PDF content for chat, length={len(document_content_for_ai)}")
                                else:
                                    # Text file
                                    with open(file_path, 'r', encoding='utf-8') as f:
                                        document_content_for_ai = f.read()
                                    logging.info(f"Loaded text content from file for chat, length={len(document_content_for_ai)}")
                            except Exception as e:
                                logging.error(f"Failed to read file for chat fallback: {e}")

                if not document_content_for_ai:
                    logging.error(f"Document {selected_document} has no readable content")
                    raise HTTPException(status_code=400, detail="Document has no analyzable content")

                logging.info(f"Found document with content length: {len(document_content_for_ai)}")

            except HTTPException:
                raise
            except Exception as e:
                logging.error(f"Error accessing document {selected_document}: {e}")
                raise HTTPException(status_code=500, detail="Failed to access document")
        
        # Build system message
        if mode == "document" and selected_document and document_content_for_ai:
            system_message = (
                "You are a legal AI assistant analyzing a specific document. "
                "Answer questions based ONLY on the content of this document. "
                "Reference specific sections when possible. "
                "If information is not found in the document, clearly state that. "
                "Do not provide general legal advice beyond what is in the document.\n\n"
                "DOCUMENT CONTENT FOR REFERENCE:\n"
                "---START OF DOCUMENT---\n"
                f"{document_content_for_ai}\n"
                "---END OF DOCUMENT---\n\n"
            )
        else:
            system_message = (
                "You are a legal AI assistant. Provide helpful, accurate legal information "
                "and analysis. Note: This is general legal information and not formal legal "
                "advice. Always consult with a qualified attorney for your specific situation. "
                "Keep responses concise and direct."
            )

        messages = [{"role": "system", "content": system_message}]

        # Add recent context (last 5 messages)
        for msg in context[-5:]:
            messages.append({"role": msg['role'], "content": msg['content']})

        messages.append({"role": "user", "content": message})

        # Get AI response
        if mode == "document" and selected_document and document_content_for_ai:
            # Create a more focused prompt for document-specific questions
            doc_prompt = f"""Based on the document content provided in your system message, answer the following question.
Answer ONLY based on what is actually in the document. If the answer is not found in the document, clearly state that.

USER QUESTION: {message}

Remember: Base your entire answer solely on the document content. Quote relevant sections when applicable."""
            
            response = await send_message(doc_prompt, mode="document", context=document_content_for_ai)
        else:
            response = await send_message(message)
            
        if not response or len(response.strip()) < 10:
            raise Exception("Empty or invalid response from AI service")
        # Strip markdown and normalize spacing
        response = strip_markdown(response)

        # Limit to approximately 100 words and make it crisp
        words = response.split()
        if len(words) > 100:
            response = ' '.join(words[:100]) + '...'
        # Make it more concise by removing unnecessary phrases
        response = re.sub(r'\s+', ' ', response)  # Multiple spaces to single
        response = response.strip()

        # Save to global chat history for all users (authenticated and anonymous)
        session_id = data.get('session_id', f"global_{uuid.uuid4()}")
        
        # For document mode, store the relevant context
        document_metadata = None
        if mode == "document" and selected_document:
            document = sync_db.documents.find_one({"id": selected_document})
            if document:
                document_metadata = {
                    "filename": document.get("filename"),
                    "summary": document.get("summary", "")[:1000],  # Store a snippet of the summary
                    "content_length": len(document.get("content", "") or document.get("document_text", ""))
                }
        
        global_chat_message = GlobalChatMessage(
            session_id=session_id,
            user_id=user_id,
            chat_mode=mode,
            selected_document=selected_document,
            question=message,
            answer=response,
            document_metadata=document_metadata
        )

        global_chat_dict = global_chat_message.model_dump()
        global_chat_dict['timestamp'] = global_chat_dict['timestamp'].isoformat()

        # Save to chat history using the existing client
        sync_db.global_chat_messages.insert_one(global_chat_dict)

        return {"response": response, "session_id": session_id}

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"❌ Chat error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Chat failed")
    finally:
        if sync_client:
            try:
                sync_client.close()
            except Exception:
                pass

@api_router.post("/tts")
async def text_to_speech(data: dict = Body(...)):
    """Converts text to speech and returns an audio file."""
    try:
        text = data.get("text")
        if not text:
            raise HTTPException(status_code=400, detail="Text is required")

        # gTTS is optional; if it's not installed, return a clear error so the server
        # can still run in environments without that dependency.
        if gTTS is None:
            raise HTTPException(status_code=501, detail="Text-to-speech not available: gTTS dependency not installed")

        # Create gTTS object
        tts = gTTS(text=text, lang='en', slow=False)

        # Save to a in-memory file
        mp3_fp = io.BytesIO()
        tts.write_to_fp(mp3_fp)
        mp3_fp.seek(0)

        return StreamingResponse(mp3_fp, media_type="audio/mpeg")

    except Exception as e:
        logging.error(f"TTS error: {e}")
        raise HTTPException(status_code=500, detail="Text-to-speech conversion failed")

@api_router.get("/analytics")
async def get_analytics_endpoint(request: Request):
    """Get user analytics"""
    try:
        # Extract token from Authorization header
        auth_header = request.headers.get("Authorization")
        if not auth_header or not auth_header.startswith("Bearer "):
            raise HTTPException(status_code=401, detail="Authorization header missing or invalid")

        token = auth_header.replace("Bearer ", "")
        payload = verify_token(token)
        if not payload:
            raise HTTPException(status_code=401, detail="Invalid token")

        user_id = payload.get("user_id")
        if not user_id:
            raise HTTPException(status_code=401, detail="Invalid token payload")

        # Get analytics data
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]

        # Get total documents
        total_docs = sync_db.documents.count_documents({"user_id": user_id})

        # Get documents from this month
        current_month = datetime.now().strftime('%Y-%m')
        this_month_count = sync_db.documents.count_documents({
            "user_id": user_id,
            "created_at": {"$gte": datetime.now().replace(day=1, hour=0, minute=0, second=0, microsecond=0)}
        })

        # Get average risk score and other analytics
        documents = list(sync_db.documents.find({"user_id": user_id}))
        risk_scores = []
        high_risk_count = 0
        risk_distribution = {'low': 0, 'medium': 0, 'high': 0, 'critical': 0}

        for doc in documents:
            try:
                analysis = json.loads(doc.get('summary', '{}'))
                score = analysis.get('riskScore', {}).get('score', 0)
                risk_scores.append(score)

                if score >= 7:
                    risk_distribution['critical'] += 1
                    high_risk_count += 1
                elif score >= 5:
                    risk_distribution['high'] += 1
                    high_risk_count += 1
                elif score >= 3:
                    risk_distribution['medium'] += 1
                else:
                    risk_distribution['low'] += 1
            except:
                pass

        avg_risk = sum(risk_scores) / len(risk_scores) if risk_scores else 0

        # Get monthly trends (last 6 months)
        monthly_trends = []
        for i in range(5, -1, -1):
            month_date = datetime.now().replace(day=1) - timedelta(days=i*30)
            month_str = month_date.strftime('%Y-%m')
            count = sync_db.documents.count_documents({
                "user_id": user_id,
                "created_at": {"$gte": month_date, "$lt": month_date + timedelta(days=30)}
            })
            monthly_trends.append({'month': month_date.strftime('%b %Y'), 'count': count})

        # Get top issues
        top_issues = []
        issue_counts = {}
        for doc in documents:
            try:
                analysis = json.loads(doc.get('summary', '{}'))
                flags = analysis.get('criticalFlags', [])
                for flag in flags:
                    title = flag.get('title', 'Unknown issue')
                    issue_counts[title] = issue_counts.get(title, 0) + 1
            except:
                pass

        top_issues = [{'issue': issue, 'count': count} for issue, count in sorted(issue_counts.items(), key=lambda x: x[1], reverse=True)[:5]]

        # Get recent activity
        recent_docs = list(sync_db.documents.find({"user_id": user_id}).sort("created_at", -1).limit(5))
        recent_activity = []
        for doc in recent_docs:
            try:
                analysis = json.loads(doc.get('summary', '{}'))
                risk_score = analysis.get('riskScore', {}).get('score', 0)
                issues_count = len(analysis.get('criticalFlags', []))
                recent_activity.append({
                    'id': doc['id'],
                    'created_at': doc.get('created_at', '').isoformat() if isinstance(doc.get('created_at'), datetime) else str(doc.get('created_at', '')),
                    'risk_score': risk_score,
                    'issues_count': issues_count
                })
            except:
                pass

        sync_client.close()

        return {
            'totalDocuments': total_docs,
            'thisMonthCount': this_month_count,
            'averageRiskScore': avg_risk,
            'highRiskCount': high_risk_count,
            'riskDistribution': risk_distribution,
            'monthlyTrends': monthly_trends,
            'topIssues': top_issues,
            'recentActivity': recent_activity
        }

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Analytics error: {e}")
        raise HTTPException(status_code=500, detail="Failed to get analytics")

@api_router.post("/compare")
async def compare_documents_endpoint(request: Request, document1: UploadFile = File(None), document2: UploadFile = File(None)):
    """Compare two documents.

    This endpoint now allows anonymous requests (no Authorization required) so users can compare two uploaded documents without signing in.
    Currently returns a simple placeholder response; file upload parameters are accepted for future processing.
    """
    try:
        # Log incoming files (if any) for debugging
        try:
            files_info = []
            if document1:
                files_info.append(f"document1={document1.filename}")
            if document2:
                files_info.append(f"document2={document2.filename}")
            logging.info(f"Compare request received. Files: {', '.join(files_info) if files_info else 'none'}")
        except Exception:
            logging.debug("No multipart files parsed or error reading file info")

        # If files were uploaded, extract text from them
        def extract_text_from_bytes(contents: bytes, filename: str, content_type: Optional[str] = None) -> str:
            text = ""
            try:
                # PDF
                if (content_type and 'pdf' in content_type) or filename.lower().endswith('.pdf'):
                    try:
                        reader = PdfReader(BytesIO(contents))
                        pages = []
                        for page in reader.pages:
                            page_text = page.extract_text()
                            if page_text and page_text.strip():
                                pages.append(page_text)
                        text = "\n".join(pages)
                    except Exception as e:
                        logging.error(f"PDF extraction error for {filename}: {e}")
                        text = ""
                else:
                    # Try decoding as utf-8, fall back to latin-1
                    try:
                        text = contents.decode('utf-8')
                    except Exception:
                        try:
                            text = contents.decode('latin-1')
                        except Exception:
                            text = ""
            except Exception as e:
                logging.error(f"Error extracting text from {filename}: {e}")
                text = ""

            return text or ""

        contents1 = None
        contents2 = None
        text1 = ""
        text2 = ""

        if document1:
            try:
                contents1 = await document1.read()
                text1 = extract_text_from_bytes(contents1, document1.filename, getattr(document1, 'content_type', None))
            except Exception as e:
                logging.error(f"Error reading document1: {e}")

        if document2:
            try:
                contents2 = await document2.read()
                text2 = extract_text_from_bytes(contents2, document2.filename, getattr(document2, 'content_type', None))
            except Exception as e:
                logging.error(f"Error reading document2: {e}")

        # If no uploaded files or no text extracted, return an informative error
        if not text1.strip() and not text2.strip():
            raise HTTPException(status_code=400, detail="No text could be extracted from the uploaded documents")

        # Compute similarity
        try:
            matcher = difflib.SequenceMatcher(None, text1, text2)
            ratio = matcher.ratio()
            similarity_percent = round(ratio * 100, 1)

            # Extract matching blocks (top matches)
            matches = []
            for block in sorted(matcher.get_matching_blocks(), key=lambda b: b.size, reverse=True):
                if block.size and block.size > 40:
                    snippet = text1[block.a:block.a + block.size].strip()
                    snippet = ' '.join(snippet.split())
                    if snippet and snippet not in matches:
                        matches.append(snippet[:500])
                if len(matches) >= 5:
                    break

            # Compute simple differences by line diff
            lines1 = [l.strip() for l in text1.splitlines() if l.strip()]
            lines2 = [l.strip() for l in text2.splitlines() if l.strip()]
            raw_diff = list(difflib.ndiff(lines1, lines2))
            diffs = []
            for d in raw_diff:
                if d.startswith('- '):
                    diffs.append(d[2:].strip())
                elif d.startswith('+ '):
                    diffs.append(d[2:].strip())
                if len(diffs) >= 10:
                    break

            # Risk heuristics: higher similarity -> lower risk (simple heuristic)
            risk1 = max(1, min(10, int(round((1 - ratio) * 10))))
            risk2 = risk1

            # Recommendations based on similarity
            if ratio >= 0.85:
                recommendations = f"Documents are highly similar ({similarity_percent}% match). Verify authorship and version control; minor wording differences may be acceptable."
            elif ratio >= 0.5:
                recommendations = f"Documents show moderate similarity ({similarity_percent}%). Review sections with differences and confirm intended variations."
            else:
                recommendations = f"Documents appear substantially different ({similarity_percent}% similarity). Perform a detailed legal review for both documents."

            return {
                "similarities": matches or [f"Documents similarity: {similarity_percent}%"],
                "differences": diffs or ["No granular line-level differences could be extracted"],
                "doc1Risk": risk1,
                "doc2Risk": risk2,
                "recommendations": recommendations,
                "similarity_percent": similarity_percent
            }
        except Exception as e:
            logging.error(f"Error computing comparison: {e}")
            raise HTTPException(status_code=500, detail="Failed to compare documents")

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Compare documents error: {e}")
        raise HTTPException(status_code=500, detail="Comparison failed")

@api_router.get("/auth/me")
async def get_current_user(request: Request):
    """Get current authenticated user info"""
    try:
        # Extract token from Authorization header
        auth_header = request.headers.get("Authorization")
        if not auth_header or not auth_header.startswith("Bearer "):
            raise HTTPException(status_code=401, detail="Authorization header missing or invalid")

        token = auth_header.replace("Bearer ", "")
        payload = verify_token(token)
        if not payload:
            raise HTTPException(status_code=401, detail="Invalid token")

        user_id = payload.get("user_id")
        if not user_id:
            raise HTTPException(status_code=401, detail="Invalid token payload")

        # Get user from database - use ObjectId for _id lookup
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
        try:
            user = sync_db.users.find_one({"_id": ObjectId(user_id)})
        except:
            user = sync_db.users.find_one({"id": user_id})
        sync_client.close()

        if not user:
            raise HTTPException(status_code=404, detail="User not found")

        return {
            "id": str(user.get("_id", user.get("id", ""))),
            "name": user.get("name", ""),
            "email": user.get("email", "")
        }

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Get current user error: {e}")
        raise HTTPException(status_code=500, detail="Failed to get user info")


@api_router.patch("/auth/me")
async def update_current_user(request: Request, data: dict = Body(...)):
    """Update current authenticated user's profile fields: name, description, avatar (base64 string).

    Expects Authorization: Bearer <token> header. Returns updated user object.
    """
    try:
        auth_header = request.headers.get("Authorization")
        if not auth_header or not auth_header.startswith("Bearer "):
            raise HTTPException(status_code=401, detail="Authorization header missing or invalid")

        token = auth_header.replace("Bearer ", "")
        payload = verify_token(token)
        if not payload:
            raise HTTPException(status_code=401, detail="Invalid token")

        user_id = payload.get("user_id")
        if not user_id:
            raise HTTPException(status_code=401, detail="Invalid token payload")

        update_fields = {}
        if 'name' in data and isinstance(data.get('name'), str):
            update_fields['name'] = data.get('name').strip()
        if 'description' in data and isinstance(data.get('description'), str):
            update_fields['description'] = data.get('description').strip()
        if 'avatar' in data and isinstance(data.get('avatar'), str):
            # store base64 data URL or any string the client sends
            update_fields['avatar'] = data.get('avatar')

        if not update_fields:
            raise HTTPException(status_code=400, detail="No valid fields to update")

        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]

        try:
            result = sync_db.users.update_one({"_id": ObjectId(user_id)}, {"$set": update_fields})
        except:
            result = sync_db.users.update_one({"id": user_id}, {"$set": update_fields})
        
        if result.matched_count == 0:
            sync_client.close()
            raise HTTPException(status_code=404, detail="User not found")

        try:
            user = sync_db.users.find_one({"_id": ObjectId(user_id)})
        except:
            user = sync_db.users.find_one({"id": user_id})
        sync_client.close()

        # Normalize returned user
        resp = {
            "id": str(user.get("_id", user.get("id", ""))), 
            "name": user.get('name'), 
            "email": user.get('email'), 
            "description": user.get('description', ''), 
            "avatar": user.get('avatar', None)
        }
        return resp

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Update current user error: {e}")
        raise HTTPException(status_code=500, detail="Failed to update user profile")

@api_router.get("/auth/google")
async def google_auth():
    """Initiate Google OAuth flow"""
    if not GOOGLE_CLIENT_ID:
        raise HTTPException(status_code=500, detail="Google OAuth not configured")

    state = secrets.token_urlsafe(32)
    google_auth_url = (
        "https://accounts.google.com/o/oauth2/v2/auth?"
        f"client_id={GOOGLE_CLIENT_ID}&"
        "response_type=code&"
        f"scope=email%20profile&"
        f"redirect_uri={GOOGLE_REDIRECT_URI}&"
        f"state={state}&"
        "access_type=offline"
    )

    return {"auth_url": google_auth_url, "state": state}

@api_router.post("/auth/google/callback", response_model=TokenResponse)
async def google_auth_callback(code: str = Form(...), state: str = Form(...)):
    """Handle Google OAuth callback"""
    try:
        if not GOOGLE_CLIENT_ID or not GOOGLE_CLIENT_SECRET:
            raise HTTPException(status_code=500, detail="Google OAuth not configured")

        # Exchange code for access token
        token_url = "https://oauth2.googleapis.com/token"
        token_data = {
            "client_id": GOOGLE_CLIENT_ID,
            "client_secret": GOOGLE_CLIENT_SECRET,
            "code": code,
            "grant_type": "authorization_code",
            "redirect_uri": GOOGLE_REDIRECT_URI,
        }

        token_response = requests.post(token_url, data=token_data)
        if token_response.status_code != 200:
            raise HTTPException(status_code=400, detail="Failed to get access token from Google")

        token_json = token_response.json()
        access_token = token_json.get("access_token")

        # Get user info from Google
        user_info_url = "https://www.googleapis.com/oauth2/v2/userinfo"
        user_info_response = requests.get(
            user_info_url,
            headers={"Authorization": f"Bearer {access_token}"}
        )

        if user_info_response.status_code != 200:
            raise HTTPException(status_code=400, detail="Failed to get user info from Google")

        user_info = user_info_response.json()
        google_id = user_info.get("id")
        email = user_info.get("email")
        name = user_info.get("name")

        # Check if user exists, create if not
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]

        user = sync_db.users.find_one({"$or": [{"email": email}, {"google_id": google_id}]})

        if not user:
            # Create new user
            user = User(
                email=email,
                name=name,
                google_id=google_id
            )
            user_dict = user.model_dump()
            user_dict['created_at'] = user_dict['created_at'].isoformat()
            sync_db.users.insert_one(user_dict)
        else:
            # Update existing user with Google ID if not present
            if not user.get('google_id'):
                sync_db.users.update_one(
                    {"_id": user["_id"]},
                    {"$set": {"google_id": google_id}}
                )

        sync_client.close()

        # Create JWT token
        jwt_token = create_access_token(
            data={"sub": email, "user_id": user['id'] if 'id' in user else user_dict['id']},
            expires_delta=timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
        )

        return TokenResponse(
            access_token=jwt_token,
            user={"id": user.get('id', user_dict['id']), "email": email, "name": name}
        )

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Google OAuth callback error: {e}")
        raise HTTPException(status_code=500, detail="OAuth authentication failed")


# Export user data (documents + chats) as a zip
@api_router.get('/auth/export')
async def export_user_data(request: Request):
    auth_header = request.headers.get('Authorization')
    token = None
    user_id = None
    if auth_header and auth_header.startswith('Bearer '):
        token = auth_header.replace('Bearer ', '')
        payload = verify_token(token)
        if payload:
            user_id = payload.get('user_id')

    # If no authenticated user, return guest history as zip if exists
    try:
        import zipfile
        buffer = BytesIO()
        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            if user_id:
                sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
                sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
                docs = list(sync_db.documents.find({'user_id': user_id}))
                chats = list(sync_db.chat_messages.find({'user_id': user_id}))
                sync_client.close()

                zf.writestr('documents.json', json.dumps(docs, default=str))
                zf.writestr('chats.json', json.dumps(chats, default=str))
            else:
                # Guest: include local server-side guest documents if any
                sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
                sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
                docs = list(sync_db.documents.find({'guest': True}).limit(50)) if 'documents' in sync_db.list_collection_names() else []
                sync_client.close()
                zf.writestr('guest_documents.json', json.dumps(docs, default=str))

        buffer.seek(0)
        return StreamingResponse(buffer, media_type='application/zip', headers={"Content-Disposition": "attachment; filename=covenantai-export.zip"})
    except Exception as e:
        logging.error(f"Export failed: {e}")
        raise HTTPException(status_code=500, detail='Export failed')


# Request password reset (generates a token and logs it / emails if configured)
@api_router.post('/auth/forgot-password')
async def forgot_password(payload: dict = Body(...)):
    email = payload.get('email')
    if not email:
        raise HTTPException(status_code=400, detail='Email is required')
    sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
    sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
    user = sync_db.users.find_one({'email': email})
    if not user:
        sync_client.close()
        # Do not reveal whether the email exists
        return {'message': 'If an account exists we sent instructions to the email provided.'}

    token = secrets.token_urlsafe(20)
    expire = (datetime.utcnow() + timedelta(hours=1)).isoformat()
    sync_db.users.update_one({'email': email}, {'$set': {'pw_reset_token': token, 'pw_reset_expires': expire}})
    sync_client.close()

    # Log and send via Resend if configured, otherwise log the token
    logging.info(f"Password reset token for {email}: {token}")
    try:
        from_email = os.environ.get('SMTP_FROM', 'CovenantAI <no-reply@waterbears.in>')
        if os.environ.get('RESEND_API_KEY'):
            subject = 'Password reset instructions'
            body = f"""
            Hello,

            Use this token to reset your password: {token}

            This token expires in 1 hour.

            If you did not request this, ignore this email.
            """
            sent = send_email_via_resend(email, subject, body, body, from_email=from_email)
            if sent:
                logging.info(f"Password reset email sent via Resend to {email}")
            else:
                logging.error(f"Resend failed to send password reset to {email}")
        else:
            logging.info('Resend not configured; password reset token logged only')
    except Exception as e:
        logging.warning('Failed to send password reset email: %s', e)

    return {'message': 'If an account exists we sent instructions to the email provided.'}


# Reset password using token
@api_router.post('/auth/reset-password')
async def reset_password(payload: dict = Body(...)):
    token = payload.get('token')
    new_password = payload.get('password')
    if not token or not new_password:
        raise HTTPException(status_code=400, detail='Token and new password are required')
    try:
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
        user = sync_db.users.find_one({'pw_reset_token': token})
        if not user:
            sync_client.close()
            raise HTTPException(status_code=400, detail='Invalid or expired token')
        expires = user.get('pw_reset_expires')
        if not expires or datetime.fromisoformat(expires) < datetime.utcnow():
            sync_client.close()
            raise HTTPException(status_code=400, detail='Token expired')

        hashed = hash_password(new_password)
        sync_db.users.update_one({'_id': user['_id']}, {'$set': {'password': hashed}, '$unset': {'pw_reset_token': '', 'pw_reset_expires': ''}})
        sync_client.close()
        return {'message': 'Password reset successful'}
    except Exception as e:
        logging.error(f"Reset password error: {e}")
        raise HTTPException(status_code=500, detail='Failed to reset password')


# Email verification (generate code)
@api_router.post('/auth/send-verification')
async def send_verification(payload: dict = Body(...)):
    email = payload.get('email')
    password = payload.get('password')
    name = payload.get('name')
    
    logging.info(f"🔵 /auth/send-verification called for {email}")
    logging.info(f"   Password: {password[:5] if password else 'NONE'}...")
    logging.info(f"   Name: {name}")
    
    if not email:
        raise HTTPException(status_code=400, detail='Email required')
    try:
        # Log SMTP configuration (excluding password)
        smtp_host = os.environ.get('SMTP_HOST')
        smtp_port = os.environ.get('SMTP_PORT', '587')
        smtp_user = os.environ.get('SMTP_USER')
        smtp_pass = os.environ.get('SMTP_PASS')
        from_email = os.environ.get('SMTP_FROM')
        
        smtp_config = {
            'SMTP_HOST': smtp_host,
            'SMTP_PORT': smtp_port,
            'SMTP_USER': smtp_user,
            'SMTP_FROM': from_email,
            'SMTP_CONFIGURED': bool(smtp_host and smtp_user and smtp_pass)
        }
        logging.info(f"SMTP Configuration: {smtp_config}")
        
        code = secrets.token_hex(3)  # 6 character verification code
        logging.info(f"Generated verification code for {email}")
        
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
        
        # Check if user exists
        user = sync_db.users.find_one({'email': email})
        
        if not user:
            # User doesn't exist - create them if password is provided (signup flow)
            if not password:
                logging.warning(f"Attempted to send verification code to non-existent user without password: {email}")
                sync_client.close()
                # Don't reveal that the user doesn't exist
                return {'message': 'If an account exists with this email, a verification code has been sent.'}
            
            # Create new user with verification code
            logging.info(f"🟢 Creating new user for {email}")
            hashed_pw = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
            logging.info(f"   Hashed password: {hashed_pw[:20]}...")
            new_user = {
                "email": email,
                "name": name or "",
                "password": hashed_pw,
                "email_verified": False,
                "email_verify_code": code,
                "email_verify_expires": (datetime.utcnow() + timedelta(minutes=15)).isoformat(),
                "created_at": datetime.utcnow().isoformat(),
            }
            logging.info(f"   New user dict: {new_user}")
            inserted = sync_db.users.insert_one(new_user)
            logging.info(f"✅ Created new user for email {email} with verification code: {code}")
            logging.info(f"   Inserted ID: {inserted.inserted_id}")
            logging.info(f"   Email verified: False")
            logging.info(f"   Code expires: {new_user['email_verify_expires']}")
        else:
            # Update existing user with new verification code
            sync_db.users.update_one({'email': email}, {
                '$set': {
                    'email_verify_code': code,
                    'email_verify_expires': (datetime.utcnow() + timedelta(minutes=15)).isoformat()
                }
            }, upsert=False)
        sync_client.close()
        logging.info(f"Updated verification code in database for {email}")

        # Prepare email settings
        smtp_host = os.environ.get('SMTP_HOST')
        smtp_port = int(os.environ.get('SMTP_PORT', '587'))
        smtp_user = os.environ.get('SMTP_USER')
        smtp_pass = os.environ.get('SMTP_PASS')
        from_email = os.environ.get('SMTP_FROM', smtp_user or 'no-reply@waterbears.in')

        subject = "Your CovenantAI Verification Code"

        body = f"""
        <html>
        <body style="margin:0; padding:0; background-color:#000; color:#fff; font-family:'Segoe UI', Roboto, Helvetica, Arial, sans-serif; text-align:center;">
            <div style="max-width:600px; margin:50px auto; padding:40px; border:1px solid #333; border-radius:12px; background-color:#0a0a0a; box-shadow:0 0 20px rgba(255,255,255,0.05);">
            
            <h1 style="font-size:28px; letter-spacing:2px; text-transform:uppercase; margin-bottom:20px; color:#ffffff;">Covenant<span style="color:#888;">AI</span></h1>
            
            <p style="font-size:16px; color:#ccc; margin-bottom:30px;">Your access code has been generated.</p>
            
            <div style="font-size:32px; letter-spacing:4px; color:#fff; background:#111; border:1px solid #333; padding:15px 0; border-radius:8px; width:80%; margin:0 auto 30px auto;">
                <strong>{code}</strong>
            </div>
            
            <p style="font-size:14px; color:#888; margin-bottom:30px;">
                This code will expire in <strong>15 minutes</strong>.<br>
                If you didn’t request this, you can safely ignore this message.
            </p>
            
            <hr style="border: none; height: 1px; background-color: #222; margin: 40px 0;">
            
            <p style="font-size:13px; color:#555; letter-spacing:1px;">— The CovenantAI Team</p>
            </div>
        </body>
        </html>
        """

        # Send via Resend only. If Resend is not configured, show the code in logs/console.
        if os.environ.get('RESEND_API_KEY'):
            try:
                sent = send_email_via_resend(email, subject, body, body, from_email=from_email)
                if sent:
                    logging.info(f"Verification email sent via Resend to {email}")
                else:
                    logging.error(f"Resend failed to send verification to {email}")
                    print("\n==================================")
                    print(f"🔑 VERIFICATION CODE for {email}: {code}")
                    print("==================================\n")
            except Exception as e:
                logging.error(f"Resend exception sending verification: {e}")
                print("\n==================================")
                print(f"🔑 VERIFICATION CODE for {email}: {code}")
                print("==================================\n")
        else:
            # Show code when Resend is not configured
            print("\n==================================")
            print(f"🔑 VERIFICATION CODE for {email}: {code}")
            print("==================================\n")

        # Don't reveal whether email exists
        return {'message': 'If an account exists with this email, a verification code has been sent.'}
    except Exception as e:
        logging.error(f"Send verification error: {e}")
        raise HTTPException(status_code=500, detail='Failed to send verification')


# Verify email code
@api_router.post('/auth/verify')
async def verify_email(payload: dict = Body(...)):
    email = payload.get('email')
    code = payload.get('code')
    if not email or not code:
        raise HTTPException(status_code=400, detail='Email and code required')
    try:
        db = get_db(async_db=False)
        user = db.users.find_one({'email': email})
        if not user:
            raise HTTPException(status_code=400, detail='Invalid email or code')
        
        # Normalize codes: strip whitespace and convert to lowercase for case-insensitive comparison
        stored_code = user.get('email_verify_code', '').strip().lower()
        input_code = code.strip().lower()
        
        # Check if code matches and hasn't expired
        if stored_code != input_code or datetime.fromisoformat(user.get('email_verify_expires', '1970-01-01T00:00:00')) < datetime.utcnow():
            raise HTTPException(status_code=400, detail='Invalid or expired code')
        db.users.update_one({'_id': user['_id']}, {'$set': {'email_verified': True}, '$unset': {'email_verify_code': '', 'email_verify_expires': ''}})
        return {'message': 'Email verified'}
    except HTTPException as http_e:
        # Re-raise HTTP exceptions to preserve the status code
        raise http_e
    except Exception as e:
        logging.error(f"Verify email error: {e}")
        raise HTTPException(status_code=500, detail='Failed to verify email')


# Delete current authenticated user
@api_router.delete('/auth/me')
async def delete_my_account(request: Request):
    auth_header = request.headers.get('Authorization')
    if not auth_header or not auth_header.startswith('Bearer '):
        raise HTTPException(status_code=401, detail='Missing authorization')
    token = auth_header.replace('Bearer ', '')
    payload = verify_token(token)
    if not payload:
        raise HTTPException(status_code=401, detail='Invalid token')
    user_id = payload.get('user_id')
    if not user_id:
        raise HTTPException(status_code=401, detail='Invalid token payload')
    try:
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
        # Delete user docs and chats
        sync_db.documents.delete_many({'user_id': user_id})
        sync_db.chat_messages.delete_many({'user_id': user_id})
        sync_db.users.delete_one({'id': user_id})
        sync_client.close()
        return {'message': 'Account deleted'}
    except Exception as e:
        logging.error(f"Delete account error: {e}")
        raise HTTPException(status_code=500, detail='Failed to delete account')

# Document upload endpoint
@api_router.post("/documents/upload")
async def upload_document(file: UploadFile = File(...)):
    """Upload a legal document for analysis"""
    logging.info(f"Received upload request for file: {file.filename}")
    try:
        # Validate file type
        allowed_types = ['application/pdf', 'text/plain', 'application/msword',
                        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        'application/octet-stream']
        file_extension = Path(file.filename).suffix.lower()
        allowed_extensions = ['.pdf', '.doc', '.docx', '.txt']
        if file.content_type not in allowed_types and file_extension not in allowed_extensions:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file.content_type} or extension: {file_extension}. Allowed types: {', '.join(allowed_types)} and extensions: {', '.join(allowed_extensions)}")
        
        # Generate unique filename
        file_id = str(uuid.uuid4())
        file_extension = Path(file.filename).suffix
        unique_filename = f"{file_id}{file_extension}"
        file_path = UPLOAD_DIR / unique_filename
        
        # Read file content fully before writing
        contents = await file.read()
        
        # Save file
        try:
            logging.info(f"Attempting to save file to: {file_path}")
            with open(file_path, "wb") as buffer:
                buffer.write(contents)
            logging.info("File saved successfully")

            # Extract text content using enhanced extraction system
            try:
                text_content, extraction_method = extract_document_text(str(file_path), file.content_type)
                
                if extraction_method == "failed" or len(text_content.strip()) < 50:
                    raise HTTPException(
                        status_code=400,
                        detail="Could not extract meaningful text from document. Please ensure the document contains readable text content."
                    )
                
                # Create enhanced document record with content stats
                content_stats = {
                    "char_count": len(text_content),
                    "word_count": len(text_content.split()),
                    "extraction_method": extraction_method,
                    "extraction_time": datetime.now(timezone.utc).isoformat()
                }
                
                document = Document(
                    filename=file.filename,
                    file_path=str(file_path),
                    file_type=file.content_type,
                    content=text_content,
                    analysis_status="pending",
                    metadata={"content_stats": content_stats}
                )
                
                # Save to database (using synchronous approach to avoid event loop conflicts)
                document_dict = document.model_dump()
                document_dict['upload_date'] = document_dict['upload_date'].isoformat()
                document_dict['content'] = text_content  # Store the extracted text
                document_dict['analysis_status'] = 'pending'

                # Use synchronous MongoDB client for this operation
                import pymongo
                sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
                sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
                sync_db.documents.insert_one(document_dict)
                
                # Trigger document analysis in the background
                asyncio.create_task(analyze_document(document.id))
                
                sync_client.close()
                
                return {"document_id": document.id, "filename": document.filename, "status": "uploaded"}
            
            except Exception as e:
                logging.error(f"Error extracting text or saving to database: {str(e)}")
                # Clean up the file if it was created but there was an error
                if file_path.exists():
                    try:
                        file_path.unlink()
                    except:
                        pass
                raise HTTPException(status_code=500, detail="Document processing failed")
            
        except Exception as e:
            logging.error(f"Error saving file: {str(e)}")
            # Clean up the file if it was created but there was a database error
            if file_path.exists():
                try:
                    file_path.unlink()
                except:
                    pass
            raise
        
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Upload error: {str(e)}")
        raise HTTPException(status_code=500, detail="File upload failed")

# Get all documents
@api_router.get("/documents", response_model=List[Document])
async def get_documents():
    """Get all uploaded documents"""
    try:
        # Use synchronous MongoDB client to avoid event loop conflicts
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
        documents = list(sync_db.documents.find())
        sync_client.close()

        safe_docs = []
        for doc in documents:
            # Normalize required fields with safe defaults to avoid Pydantic validation errors
            safe = {}
            # id: prefer 'id' field, otherwise fallback to MongoDB _id string
            safe['id'] = doc.get('id') or str(doc.get('_id'))
            safe['filename'] = doc.get('filename', 'Untitled Document')
            safe['file_path'] = doc.get('file_path', '')
            safe['file_type'] = doc.get('file_type', 'application/octet-stream')

            # upload_date may be stored as datetime or ISO string; normalize to datetime
            ud = doc.get('upload_date')
            if isinstance(ud, datetime):
                safe['upload_date'] = ud
            elif isinstance(ud, str):
                try:
                    safe['upload_date'] = datetime.fromisoformat(ud)
                except Exception:
                    safe['upload_date'] = datetime.now(timezone.utc)
            else:
                safe['upload_date'] = datetime.now(timezone.utc)

            safe['analysis_status'] = doc.get('analysis_status', 'completed')
            safe['summary'] = doc.get('summary')
            safe['key_clauses'] = doc.get('key_clauses')
            safe['risk_assessment'] = doc.get('risk_assessment')

            safe_docs.append(safe)

        return [Document(**d) for d in safe_docs]
    except Exception as e:
        logging.error(f"Get documents error: {e}")
        raise HTTPException(status_code=500, detail="Failed to retrieve documents")

# Delete document endpoint
@api_router.delete("/documents/{document_id}")
async def delete_document(document_id: str):
    """Delete a document by ID"""
    try:
        # Log the delete request
        logging.info(f"Attempting to delete document with ID: {document_id}")

        # Use synchronous MongoDB client to avoid event loop conflicts
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]

        # Try to find by 'id' field first, then by '_id' field (MongoDB ObjectId)
        document = sync_db.documents.find_one({"id": document_id})
        if not document:
            try:
                # Try as MongoDB ObjectId
                from bson import ObjectId
                document = sync_db.documents.find_one({"_id": ObjectId(document_id)})
            except Exception:
                pass
        
        if not document:
            sync_client.close()
            logging.warning(f"Document not found for deletion: {document_id}")
            raise HTTPException(status_code=404, detail="Document not found")

        # Delete the file from the filesystem
        try:
            file_path = Path(document.get('file_path', ''))
            if file_path and file_path.exists():
                file_path.unlink()
                logging.info(f"File deleted successfully: {file_path}")
        except Exception as e:
            logging.error(f"Error deleting file: {e}")
            # Continue with document deletion even if file deletion fails

        # Delete the document record from the database - use the actual _id
        result = sync_db.documents.delete_one({"_id": document["_id"]})
        sync_client.close()

        if result.deleted_count == 0:
            logging.warning(f"Document not deleted from database: {document_id}")
            raise HTTPException(status_code=404, detail="Document not found in database")

        logging.info(f"Document deleted successfully: {document_id}")
        return {"message": "Document deleted successfully", "id": document_id}

    except HTTPException as he:
        raise he
    except Exception as e:
        logging.error(f"Delete document error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Get specific document
@api_router.get("/documents/{document_id}")
async def get_document(document_id: str):
    """Get a specific document by ID"""
    try:
        # Use synchronous MongoDB client to avoid event loop conflicts
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
        document = sync_db.documents.find_one({"id": document_id})
        sync_client.close()

        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        if isinstance(document.get('upload_date'), str):
            document['upload_date'] = datetime.fromisoformat(document['upload_date'])

        return Document(**document)
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Get document error: {e}")
        raise HTTPException(status_code=500, detail="Failed to retrieve document")




MAX_CHUNK_SIZE = 8000

def generate_analysis_prompt(text: str) -> str:
    return (
        "You are a legal document analysis expert. Analyze this document content and provide a "
        "structured analysis in JSON format.\n\n"
        f"DOCUMENT CONTENT:\n{text}\n\n"
        "Provide your response as a JSON object with these exact fields:\n"
        "{\n"
        '  "document_type": {"type": "string", "confidence": "High/Medium/Low"},\n'
        '  "key_points": [\n'
        '    {"point": "string", "importance": "Critical/High/Medium/Low"}\n'
        "  ],\n"
        '  "parties": [\n'
        '    {"name": "string", "role": "string", "obligations": ["string"]}\n'
        "  ],\n"
        '  "risks": [\n'
        '    {"risk": "string", "severity": "High/Medium/Low", "mitigation": "string"}\n'
        "  ],\n"
        '  "recommendations": [\n'
        '    {"action": "string", "priority": "High/Medium/Low", "rationale": "string"}\n'
        "  ]\n"
        "}\n\n"
        "REQUIREMENTS:\n"
        "1. Base analysis ONLY on the provided content\n"
        "2. Include specific references to document text\n"
        "3. Focus on legal implications and risks\n"
        "4. Keep explanations clear and actionable\n"
        "5. Ensure output is valid JSON"
    )

@api_router.post("/documents/{document_id}/analyze")
async def analyze_document(document_id: str, request: Request = None):
    """Analyze a document with the given ID. The request parameter is optional and only used for auth."""
    sync_client = None
    user_id = None
    all_analyses: List[Dict[str, Any]] = []
    try:
        # Extract user_id from Authorization header if present
        auth_header = request.headers.get("Authorization")
        if auth_header and auth_header.startswith("Bearer "):
            token = auth_header.replace("Bearer ", "")
            payload = verify_token(token)
            if payload:
                user_id = payload.get("user_id")
                logging.info(f"Analyze request from user: {user_id}")
        
        # Connect to DB
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]

        document = sync_db.documents.find_one({"id": document_id})
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        # Mark processing
        sync_db.documents.update_one({"id": document_id}, {"$set": {"analysis_status": "processing"}})

        # Get content or fallback to reading file
        content = document.get('content', '') or document.get('document_text', '')
        document_text = content or ""

        # If no inline content, try to read from file_path
        if not document_text or len(document_text.strip()) < 50:
            file_path = document.get('file_path')
            if file_path:
                fp = Path(file_path)
                if not fp.exists():
                    logging.warning(f"File not found at path: {file_path}")
                else:
                    # If PDF, use PyPDF2 then fallbacks
                    try:
                        if document.get('file_type') == 'application/pdf' or fp.suffix.lower() == '.pdf':
                            from PyPDF2 import PdfReader
                            reader = PdfReader(str(fp))
                            pages = []
                            for p in reader.pages:
                                txt = p.extract_text() or ''
                                pages.append(txt)
                            document_text = "\n".join(pages).strip()
                            if len(document_text) < 50:
                                pdf_text = extract_text_pdf_fallback(str(fp))
                                if pdf_text and len(pdf_text.strip()) > len(document_text):
                                    document_text = pdf_text
                                else:
                                    ocr_text = extract_text_with_ocr(str(fp))
                                    if ocr_text and len(ocr_text.strip()) > len(document_text):
                                        document_text = ocr_text
                        else:
                            # docx fallback or plain text
                            if fp.suffix.lower() == '.docx':
                                document_text = extract_text_docx(str(fp)) or document_text
                            if not document_text:
                                # try reading as utf-8
                                with open(fp, 'r', encoding='utf-8') as f:
                                    document_text = f.read()
                    except Exception as e:
                        logging.exception("Error extracting text from file")
                        # keep document_text as-is or empty

        if not document_text or len(document_text.strip()) < 50:
            sync_db.documents.update_one({"id": document_id}, {"$set": {"analysis_status": "failed", "analysis_error": "Insufficient text for analysis"}})
            raise HTTPException(status_code=400, detail="Document has insufficient content for analysis")

        logging.info(f"Starting analysis for document {document_id}, extracted length={len(document_text)}")

        # Chunk and analyze
        chunks = [document_text[i:i + MAX_CHUNK_SIZE] for i in range(0, len(document_text), MAX_CHUNK_SIZE)]
        for i, chunk in enumerate(chunks):
            logging.info(f"Analyzing chunk {i+1}/{len(chunks)}")
            prompt = generate_analysis_prompt(chunk)
            try:
                resp = await send_message(prompt)
                resp = resp.strip()
                # remove triple-backtick fences if present
                if resp.startswith(""):
                    # support json or 
                    resp = resp.split("\n", 1)[1] if "\n" in resp else resp[3:]
                    if resp.endswith(""):
                        resp = resp[:-3]
                    resp = resp.strip()
                parsed = json.loads(resp)
                all_analyses.append(parsed)
            except json.JSONDecodeError:
                logging.exception("AI returned non-JSON for a chunk; skipping that chunk")
                continue
            except Exception:
                logging.exception("Error calling send_message for chunk")
                continue

        if not all_analyses:
            sync_db.documents.update_one({"id": document_id}, {"$set": {"analysis_status": "failed", "analysis_error": "AI returned no valid JSON responses"}})
            raise HTTPException(status_code=500, detail="Failed to get valid analysis from AI")

        # Merge analyses
        merged = {
            "document_type": all_analyses[0].get("document_type", {}),
            "key_points": [],
            "risks": [],
            "recommendations": []
        }
        seen_points = set()
        seen_risks = set()
        seen_recs = set()

        for analysis in all_analyses:
            for kp in analysis.get("key_points", []):
                text = kp.get("point", "").strip()
                if text and text not in seen_points:
                    merged["key_points"].append(kp)
                    seen_points.add(text)
            for r in analysis.get("risks", []):
                text = r.get("risk", "").strip()
                if text and text not in seen_risks:
                    merged["risks"].append(r)
                    seen_risks.add(text)
            for rc in analysis.get("recommendations", []):
                text = rc.get("action", "").strip()
                if text and text not in seen_recs:
                    merged["recommendations"].append(rc)
                    seen_recs.add(text)

        # Prepare update dict
        update_dict = {
            "analysis_status": "completed",
            "document_type": merged["document_type"],
            "key_points": merged["key_points"],
            "risks": merged["risks"],
            "recommendations": merged["recommendations"],
            "analysis_version": "2.0",
            "last_analyzed": datetime.now(timezone.utc).isoformat()
        }
        
        # If user is authenticated, add user_id to mark this as user's document
        if user_id:
            update_dict["user_id"] = user_id
            logging.info(f"✅ Saving analysis with user_id: {user_id}")
        
        # Update DB with completed analysis
        sync_db.documents.update_one(
            {"id": document_id},
            {"$set": update_dict}
        )

        logging.info(f"Analysis completed for document {document_id}")
        return merged

    except HTTPException:
        # re-raise to let FastAPI return it unchanged
        raise
    except Exception as e:
        logging.exception(f"Unexpected error analyzing document {document_id}")
        # Attempt to mark failed
        try:
            if sync_client:
                sync_db.documents.update_one({"id": document_id}, {"$set": {"analysis_status": "failed", "analysis_error": str(e)}})
        except Exception:
            logging.exception("Failed to write failure status to DB")
        raise HTTPException(status_code=500, detail=f"Analysis failed: {e}")
    finally:
        if sync_client:
            try:
                sync_client.close()
            except Exception:
                logging.exception("Error closing Mongo client")

# New /api/analyze endpoint for direct text analysis and file upload
@api_router.post("/analyze")
async def analyze_text(file: UploadFile = File(None), rawText: str = Form(None)):
    """Analyze raw text or uploaded file and return structured JSON response"""
    try:
        document_text = ""

        if file:
            logging.info(f"Received file upload analysis request. Filename: {file.filename}, Content-Type: {file.content_type}")

            # Validate file type
            allowed_types = ['application/pdf', 'text/plain', 'application/msword',
                            'application/vnd.openxmlformats-officedocument.wordprocessingml.document']
            if file.content_type not in allowed_types and not file.filename.lower().endswith(('.pdf', '.doc', '.docx', '.txt')):
                raise HTTPException(status_code=400, detail=f"Unsupported file type: {file.content_type}")

            # Read file content
            contents = await file.read()

            if file.content_type == 'application/pdf' or file.filename.lower().endswith('.pdf'):
                # Extract text from PDF
                try:
                    from io import BytesIO
                    pdf_reader = PdfReader(BytesIO(contents))
                    text_pages = []
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_pages.append(page_text)
                    document_text = "\n".join(text_pages)

                    if len(document_text.strip()) < 50:
                        document_text = "This PDF appears to be image-based or scanned. OCR processing is not available. Please provide a text-based PDF or paste the document text directly."
                except Exception as e:
                    document_text = f"Failed to extract text from PDF: {str(e)}"
            else:
                # For text files
                try:
                    document_text = contents.decode('utf-8')
                except UnicodeDecodeError:
                    document_text = "Unable to decode file content. Please ensure it's a text-based document."

        if rawText:
            logging.info(f"Received direct text analysis request. Text length: {len(rawText)}")
            document_text = rawText

        if not document_text or len(document_text.strip()) < 10:
            raise HTTPException(status_code=400, detail="Document content is too short for analysis")

        # Create analysis prompt for the document text
        analysis_prompt = f"""You are a legal document analysis expert. Analyze the following legal document text and provide a comprehensive analysis in the exact JSON format specified below.

LEGAL DOCUMENT TEXT TO ANALYZE:
{document_text[:120000]}

IMPORTANT: Provide your response as a valid JSON object with these exact keys and structure:

{{
  "summary": "A clear, concise summary of what this document is about in plain English",
  "riskScore": {{
    "score": 7,
    "max": 10,
    "label": "High Risk - Requires Immediate Legal Review"
  }},
  "analysis": {{
    "strengths": [
      {{"text": "Clear and unambiguous language"}},
      {{"text": "Well-defined obligations for both parties"}},
      {{"text": "Reasonable termination clauses"}}
    ],
    "weaknesses": [
      {{"text": "Unfavorable payment terms"}},
      {{"text": "Limited liability protection"}},
      {{"text": "Ambiguous dispute resolution"}}
    ],
    "opportunities": [
      {{"text": "Negotiate better payment terms"}},
      {{"text": "Add stronger liability protections"}},
      {{"text": "Clarify dispute resolution process"}}
    ],
    "threats": [
      {{"text": "Potential financial losses"}},
      {{"text": "Legal disputes and litigation costs"}},
      {{"text": "Damage to business reputation"}}
    ]
  }},
  "criticalFlags": [
    {{
      "title": "Unfavorable Payment Terms",
      "explanation": "The payment terms heavily favor the other party with delayed payments and high penalties.",
      "source": "Section 4.2 Payment Terms"
    }},
    {{
      "title": "Weak Liability Protection",
      "explanation": "Liability limitations are insufficient and may not hold up in court.",
      "source": "Section 7.1 Limitation of Liability"
    }}
  ],
  "negotiationPoints": [
    {{
      "title": "Payment Terms",
      "risk": "High - Could impact cash flow significantly",
      "example": "Request 50% payment upfront and 50% upon completion"
    }},
    {{
      "title": "Liability Cap",
      "risk": "Medium - May expose to unlimited liability",
      "example": "Cap liability at 2x the contract value"
    }}
  ]
}}

REQUIREMENTS:
- Base ALL analysis on the provided document text
- Use the exact JSON structure shown above
- Provide realistic, specific analysis based on the document content
- Include 3-5 items in each analysis array
- Make criticalFlags and negotiationPoints specific to the document
- Ensure the response is valid JSON that can be parsed

If the document is too short or unclear, provide a general analysis framework but still use the required JSON structure."""

        # Get AI analysis
        analysis_response = await send_message(analysis_prompt)
        logging.info(f"AI analysis response received. Length: {len(analysis_response)}")

        # Parse the JSON response
        try:
            # Clean the response to ensure it's valid JSON
            analysis_response = analysis_response.strip()
            if analysis_response.startswith('```json'):
                analysis_response = analysis_response[7:]
            if analysis_response.endswith('```'):
                analysis_response = analysis_response[:-3]
            analysis_response = analysis_response.strip()

            # Try to parse JSON
            try:
                result = json.loads(analysis_response)
                logging.info("Successfully parsed analysis JSON")
            except json.JSONDecodeError as e:
                logging.error(f"JSON parsing error: {e}")
                logging.error(f"Raw response: {analysis_response[:1000]}")

                # Try to extract JSON from the response if it's wrapped in other text
                import re
                json_match = re.search(r'\{.*\}', analysis_response, re.DOTALL)
                if json_match:
                    try:
                        result = json.loads(json_match.group())
                        logging.info("Successfully extracted and parsed JSON from response")
                    except json.JSONDecodeError:
                        logging.warning("Failed to parse extracted JSON; will fallback to raw AI text instead of returning 500")
                        # Fall back to a safe result containing the raw AI response so downstream
                        # consumers (frontend/export) can still show the full analysis text.
                        result = {
                            "summary": (analysis_response[:200] + '...') if len(analysis_response) > 200 else analysis_response,
                            "riskScore": {"score": 0, "max": 10, "label": "Unknown"},
                            "analysis": {"raw_text": analysis_response},
                            "criticalFlags": [],
                            "negotiationPoints": []
                        }
                else:
                    logging.warning("AI response did not contain JSON; falling back to raw AI text")
                    result = {
                        "summary": (analysis_response[:200] + '...') if len(analysis_response) > 200 else analysis_response,
                        "riskScore": {"score": 0, "max": 10, "label": "Unknown"},
                        "analysis": {"raw_text": analysis_response},
                        "criticalFlags": [],
                        "negotiationPoints": []
                    }

            # Ensure required fields exist in the result (if we created a fallback above,
            # they will already be present). If some are missing, add safe defaults so
            # downstream code can rely on the structure.
            required_fields = ['summary', 'riskScore', 'analysis', 'criticalFlags', 'negotiationPoints']
            for field in required_fields:
                if field not in result:
                    logging.warning(f"Missing field '{field}' in AI result; inserting default value")
                    if field == 'summary':
                        result['summary'] = ''
                    elif field == 'riskScore':
                        result['riskScore'] = {"score": 0, "max": 10, "label": "Unknown"}
                    elif field == 'analysis':
                        result['analysis'] = {"raw_text": analysis_response}
                    else:
                        result[field] = []

            # Ensure SWOT keys exist and are lists so the frontend can safely render them
            analysis_obj = result.get('analysis', {}) if isinstance(result.get('analysis', {}), dict) else {}
            for k in ('strengths', 'weaknesses', 'opportunities', 'threats'):
                if k not in analysis_obj or not isinstance(analysis_obj.get(k), list):
                    analysis_obj[k] = []
            result['analysis'] = analysis_obj

            return result

        except json.JSONDecodeError as e:
            logging.error(f"JSON parsing error: {e}")
            logging.error(f"Raw response: {analysis_response[:500]}")
            raise HTTPException(status_code=500, detail=f"Failed to parse AI response as JSON: {str(e)}")

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Analysis error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")




@api_router.post("/documents/ask")
async def ask_question(request: QuestionRequest):
    """Ask a specific question about a document."""
    sync_client = None
    try:
        # --- Connect to database ---
        mongo_url = os.environ.get('MONGO_URL', 'mongodb://localhost:27017')
        db_name = os.environ.get('DB_NAME', 'legal_docs')
        sync_client = pymongo.MongoClient(mongo_url)
        sync_db = sync_client[db_name]

        # --- Get document ---
        document = sync_db.documents.find_one({"id": request.document_id})
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        # --- Prepare session ID ---
        session_id = request.session_id or f"qa_{request.document_id}_{uuid.uuid4()}"

        # --- Prepare file content wrapper ---
        file_content = FileContentWithMimeType(
            file_path=document['file_path'],
            mime_type=document['file_type']
        )

        # --- Extract text from document ---
        document_text = ""
        try:
            if document['file_type'] == 'application/pdf' or document['file_path'].lower().endswith('.pdf'):
                # Try PDF text extraction
                try:
                    reader = PdfReader(document['file_path'])
                    text_pages = [page.extract_text() or "" for page in reader.pages]
                    document_text = "\n".join([t.strip() for t in text_pages if t and t.strip()])

                    # If extracted text too short → attempt OCR
                    if len(document_text.strip()) < 50:
                        logging.info("Attempting OCR for image-based/scanned PDF")
                        ocr_text = extract_text_with_ocr(document['file_path'])
                        if len(ocr_text.strip()) > len(document_text.strip()):
                            document_text = ocr_text
                            logging.info(f"OCR improved extraction length: {len(document_text)} chars")
                        else:
                            msg = (
                                f"This PDF seems scanned or image-based. "
                                f"Text extraction yielded only {len(document_text)} chars. "
                                "OCR attempted but results insufficient. Please use a text-based PDF."
                            )
                            logging.warning(msg)
                            document_text = msg
                    else:
                        logging.info(f"Extracted {len(document_text)} chars from PDF")
                except Exception as e:
                    logging.error(f"PDF extraction error: {e}")
                    document_text = f"Failed to extract text from PDF: {str(e)}"
            else:
                # Non-PDF files
                with open(document['file_path'], 'r', encoding='utf-8') as f:
                    document_text = f.read()
                logging.info(f"Extracted {len(document_text)} chars from text file")
        except UnicodeDecodeError:
            # Retry with bytes
            try:
                with open(document['file_path'], 'rb') as f:
                    document_bytes = f.read()
                document_text = document_bytes.decode('utf-8', errors='ignore')
            except Exception:
                document_text = "Unable to decode file; likely binary format."
        except Exception as e:
            document_text = f"Error reading document: {str(e)}"

        # --- Create question prompt ---
        question_prompt = f"""
You are a legal document analysis expert.
Below is the content of a legal document. Please answer the question based ONLY on this document.

DOCUMENT CONTENT:
{document_text[:10000]}

QUESTION: {request.question}

REQUIREMENTS:
1. Base your answer solely on this document.
2. Reference specific clauses or text segments when possible.
3. Avoid generic legal advice—cite this document’s content.
4. Provide:
   - A direct answer,
   - Supporting sections,
   - Context or implications,
   - Practical advice relevant to this document.
Keep the explanation clear but accurate to the text.
        """.strip()

        # --- Send to model ---
        question_message = UserMessage(text=question_prompt)
        answer = await send_message(question_prompt)

        if not answer or len(answer.strip()) < 10:
            raise Exception("Empty or invalid response from model")

        # --- Normalize answer ---
        answer = strip_markdown(answer)
        words = answer.split()
        if len(words) > 200:
            answer = ' '.join(words[:200]) + '...'
        answer = re.sub(r'\s+', ' ', answer).strip()

        # --- Save chat message ---
        chat_message = ChatMessage(
            document_id=request.document_id,
            session_id=session_id,
            question=request.question,
            answer=answer
        )
        chat_dict = chat_message.model_dump()
        chat_dict['timestamp'] = chat_dict['timestamp'].isoformat()

        sync_db.chat_messages.insert_one(chat_dict)

        # --- Return response ---
        return {
            "question": request.question,
            "answer": answer,
            "session_id": session_id
        }

    except HTTPException:
        raise
    except Exception as e:
        logging.exception("Error in /documents/ask")
        raise HTTPException(status_code=500, detail=f"Failed to answer question: {str(e)}")
    finally:
        if sync_client:
            try:
                sync_client.close()
            except Exception:
                logging.error("Error closing MongoDB connection")


@api_router.post("/documents/ask/inline")
async def ask_inline(payload: dict = Body(...)):
    """Ask a specific question about a provided document text or analysis (no DB entry required).

    Expected payload:
    {
      "question": "...",
      "document_text": "..."  # optional if analysis provided
      "document_id": "..." # optional, will try to load from database
      "analysis": {...} # optional, used when provided (analysis may contain raw_text)
    }
    """
    try:
        question = payload.get('question')
        if not question:
            raise HTTPException(status_code=400, detail="Question is required")

        # Prefer raw text from analysis if provided
        document_text = ''
        if payload.get('analysis') and isinstance(payload.get('analysis'), dict):
            analysis = payload.get('analysis')
            document_text = analysis.get('raw_text') or analysis.get('document_text') or ''

        # Fallback to explicit document_text
        if not document_text:
            document_text = payload.get('document_text', '')

        # If still empty, try to load from database if document_id is provided
        if not document_text or len(document_text.strip()) < 20:
            document_id = payload.get('document_id')
            logging.info(f"[INLINE ASK] document_id={document_id}, document_text_len={len(document_text.strip())}")
            if document_id:
                try:
                    sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
                    sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
                    # Try to find by 'id' field first, then by '_id' field (MongoDB ObjectId)
                    document = sync_db.documents.find_one({"id": document_id})
                    logging.info(f"[INLINE ASK] Search by id: {document is not None}")
                    if not document:
                        # Try finding by _id (in case document_id is actually a MongoDB _id)
                        try:
                            document = sync_db.documents.find_one({"_id": ObjectId(document_id)})
                            logging.info(f"[INLINE ASK] Search by _id: {document is not None}")
                        except Exception as e:
                            # If it's not a valid ObjectId, just skip this search
                            logging.debug(f"Document ID {document_id} is not a valid ObjectId: {e}")
                    sync_client.close()
                    
                    if document:
                        logging.info(f"[INLINE ASK] Document found, ID={document.get('id', 'NO_ID')}")
                        # Try to get content from database fields
                        db_text = (
                            document.get('content', '') or 
                            document.get('document_text', '') or
                            document.get('summary', '')
                        )
                        logging.info(f"[INLINE ASK] DB content length: {len(db_text) if db_text else 0}")
                        
                        # If still empty, try to read from file
                        if not db_text:
                            file_path = document.get('file_path')
                            logging.info(f"[INLINE ASK] No DB content, trying file_path={file_path}")
                            if file_path:
                                try:
                                    from pathlib import Path
                                    if Path(file_path).exists():
                                        logging.info(f"[INLINE ASK] File exists: {file_path}")
                                        if file_path.lower().endswith('.pdf'):
                                            from PyPDF2 import PdfReader
                                            reader = PdfReader(file_path)
                                            pages = []
                                            for page in reader.pages:
                                                text = page.extract_text()
                                                if text:
                                                    pages.append(text)
                                            db_text = "\n".join(pages)
                                        else:
                                            with open(file_path, 'r', encoding='utf-8') as f:
                                                db_text = f.read()
                                        logging.info(f"[INLINE ASK] Loaded from file, length={len(db_text)}")
                                    else:
                                        logging.warning(f"[INLINE ASK] File does not exist: {file_path}")
                                except Exception as e:
                                    logging.error(f"Failed to read file from inline endpoint: {e}")
                        
                        if db_text:
                            document_text = db_text
                            logging.info(f"[INLINE ASK] Loaded document content, final_length={len(document_text)}")
                    else:
                        logging.info(f"[INLINE ASK] No document found with ID {document_id}")
                except Exception as e:
                    logging.error(f"Failed to load document from database in inline endpoint: {e}")

        # If we have very little context, still call the general chat endpoint
        if not document_text or len(document_text.strip()) < 20:
            # Build a short prompt describing that document context is limited
            prompt = f"You are an expert legal assistant. The user asks: {question}. Note: limited or no document content was provided. Answer concisely."
        else:
            # Build a prompt that instructs the LLM to answer based ONLY on the provided document text
            prompt = f"You are a legal document analysis expert. Use ONLY the document text provided below to answer the question. DOCUMENT TEXT:\n{document_text[:12000]}\nQUESTION: {question}\nPlease answer directly, reference the document where appropriate, and keep the answer concise."

        answer = await send_message(prompt)
        if not answer or len(answer.strip()) < 1:
            raise Exception("Empty response from AI service")
        # Strip markdown from AI answer and optionally persist the QA pair
        answer = strip_markdown(answer)
        # Normalize and truncate similar to general chat
        words = answer.split()
        if len(words) > 200:
            answer = ' '.join(words[:200]) + '...'
        answer = re.sub(r'\s+', ' ', answer).strip()

        # Optionally persist the QA pair to chat_messages for continuity (non-blocking)
        try:
            chat_message = ChatMessage(
                document_id=payload.get('document_id', f"inline_{uuid.uuid4()}"),
                session_id=payload.get('session_id', f"inline_{uuid.uuid4()}"),
                question=question,
                answer=answer
            )
            chat_dict = chat_message.model_dump()
            chat_dict['timestamp'] = chat_dict['timestamp'].isoformat()
            sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
            sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
            sync_db.chat_messages.insert_one(chat_dict)
            sync_client.close()
        except Exception:
            # Non-fatal if saving fails
            logging.debug("Failed to persist inline chat message (non-fatal)")

        return {"answer": answer}

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Inline question answering error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to answer inline question: {str(e)}")

# Get chat history for a document
@api_router.get("/documents/{document_id}/chat")
async def get_chat_history(document_id: str, session_id: Optional[str] = None):
    """Get chat history for a document"""
    try:
        query = {"document_id": document_id}
        if session_id:
            query["session_id"] = session_id

        # Use synchronous MongoDB client to avoid event loop conflicts
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
        messages = list(sync_db.chat_messages.find(query).sort("timestamp", 1))
        sync_client.close()

        for msg in messages:
            if isinstance(msg.get('timestamp'), str):
                msg['timestamp'] = datetime.fromisoformat(msg['timestamp'])

        return [ChatMessage(**msg) for msg in messages]

    except Exception as e:
        logging.error(f"Get chat history error: {e}")
        raise HTTPException(status_code=500, detail="Failed to retrieve chat history")



# Export selective sections of document analysis as PDF
@api_router.post("/documents/{document_id}/export-selective-pdf")
async def export_selective_pdf(document_id: str, request_body: ExportRequest):
    """Generate a sleek, minimal black-and-white PDF with selected sections."""
    logging.info(f"Selective PDF export started for document: {document_id}")

    try:
        sections = request_body.sections or []
        if not sections:
            raise HTTPException(status_code=400, detail="No sections selected for export")

        # --- Fetch document from DB ---
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
        document = db.documents.find_one({"id": document_id})
        sync_client.close()

        if not document:
            logging.error(f"Document not found: {document_id}")
            raise HTTPException(status_code=404, detail="Document not found")

        if document.get('analysis_status') != 'completed':
            logging.error(f"Document analysis incomplete: {document_id}")
            raise HTTPException(status_code=400, detail="Document analysis not completed yet")

        # --- Helper: Clean text for PDF ---
        def clean_text(text: str) -> str:
            if not text:
                return ""
            t = str(text)
            t = t.replace('\x00', '').replace('•', '*').replace('\u2022', '*')
            t = t.replace('–', '-').replace('—', '-')
            t = t.replace('\u2018', "'").replace('\u2019', "'").replace('\u201C', '"').replace('\u201D', '"')
            return t.strip()

        # --- Setup FPDF ---
        from fpdf import FPDF
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)

        try:
            dejavu = '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'
            if os.path.exists(dejavu):
                pdf.add_font('DejaVu', '', dejavu, uni=True)
                pdf.add_font('DejaVu', 'B', dejavu, uni=True)
                base_font = 'DejaVu'
            else:
                base_font = 'Arial'
        except Exception:
            base_font = 'Arial'

        pdf.add_page()
        pdf.set_font(base_font, 'B', 16)
        pdf.cell(0, 10, "Legal Document Analysis Report", ln=True, align='C')
        pdf.ln(6)

        pdf.set_font(base_font, size=11)
        pdf.cell(0, 8, f"Document: {document.get('filename', 'Unknown')}", ln=True)
        pdf.cell(0, 8, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True)
        pdf.ln(6)

        # --- Section: Summary ---
        if 'summary' in sections:
            summary = document.get('executive_summary') or ''
            plain_english = document.get('plain_english') or ''
            if summary or plain_english:
                pdf.set_font(base_font, 'B', 12)
                pdf.cell(0, 8, "Summary", ln=True)
                pdf.set_font(base_font, size=10)
                pdf.multi_cell(0, 5, clean_text(summary)[:3000] + ("..." if len(summary) > 3000 else ""))
                pdf.multi_cell(0, 5, clean_text(plain_english))
                pdf.ln(4)

        # --- Section: Key Clauses ---
        if 'keyclauses' in sections:
            clauses = document.get('key_clauses') or []
            if clauses:
                pdf.set_font(base_font, 'B', 12)
                pdf.cell(0, 8, "Key Clauses", ln=True)
                for clause in clauses:
                    if clause.get('clause'):
                        pdf.set_font(base_font, 'B', 10)
                        pdf.cell(0, 6, f"Clause: {clean_text(clause['clause'][:100])}", ln=True)
                        if clause.get('explanation'):
                            pdf.set_font(base_font, size=9)
                            pdf.multi_cell(0, 4, clean_text(clause['explanation']))
                        pdf.ln(2)

        # --- Section: Risk Assessment ---
        if 'risks' in sections:
            risks = document.get('risk_assessment') or ''
            if risks:
                pdf.set_font(base_font, 'B', 12)
                pdf.cell(0, 8, "Risk Assessment", ln=True)
                pdf.set_font(base_font, size=10)
                pdf.multi_cell(0, 5, clean_text(risks))
                pdf.ln(3)

        # --- Section: Q&A ---
        if 'qa' in sections:
            sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
            qa_list = list(sync_client[os.environ.get('DB_NAME', 'legal_docs')]
                           .chat_messages.find({"document_id": document_id}).limit(10))
            sync_client.close()

            if qa_list:
                pdf.set_font(base_font, 'B', 12)
                pdf.cell(0, 8, "Questions & Answers", ln=True)
                pdf.set_font(base_font, size=10)
                for qa in qa_list:
                    q = clean_text(qa.get('question', ''))
                    a = clean_text(qa.get('answer', ''))
                    pdf.set_font(base_font, 'B', 10)
                    pdf.cell(0, 6, f"Q: {q[:100]}", ln=True)
                    pdf.set_font(base_font, size=9)
                    pdf.multi_cell(0, 4, f"A: {a}")
                    pdf.ln(2)

        # --- Finalize and return PDF ---
        tmp_file = tempfile.mktemp(suffix='.pdf')
        pdf.output(tmp_file)

        filename = f"legal_analysis_{document_id}.pdf"
        logging.info(f"Selective PDF generated: {filename}")

        return FileResponse(
            tmp_file,
            media_type="application/pdf",
            filename=filename,
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"PDF export failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")




# Export report data as PDF (CovenantAI Style: Black & White, Sleek & Symmetric)
@api_router.post("/export/pdf")
async def export_report_pdf(report_data: dict = Body(...), bundle: Optional[bool] = False):
    """Export CovenantAI report as a minimal, symmetric, black-and-white styled PDF"""
    logging.info("Starting CovenantAI PDF export")

    try:
        # --- Helper: Clean text ---
        def clean_text_for_pdf(text: str) -> str:
            if not text:
                return ""
            text = str(text)
            text = text.replace('\x00', '')
            text = (
                text.replace('•', '*')
                    .replace('–', '-')
                    .replace('—', '-')
                    .replace('\u2018', "'")
                    .replace('\u2019', "'")
                    .replace('\u201C', '"')
                    .replace('\u201D', '"')
            )
            return text.strip()

        from fpdf import FPDF
        import tempfile, json, os, re
        from datetime import datetime

        # --- Setup PDF ---
        pdf = FPDF(format='A4')
        try:
            dejavu_path = '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'
            if os.path.exists(dejavu_path):
                pdf.add_font('DejaVu', '', dejavu_path, uni=True)
                pdf.add_font('DejaVu', 'B', dejavu_path, uni=True)
                base_font = 'DejaVu'
            else:
                base_font = 'Arial'
        except Exception:
            base_font = 'Arial'

        logging.info(f"Selected base_font={base_font}")

        pdf.add_page()
        pdf.set_left_margin(15)
        pdf.set_right_margin(15)

        # --- Header ---
        pdf.set_font(base_font, 'B', 16)
        pdf.cell(0, 10, "COVENANTAI REPORT", ln=True, align='C')
        pdf.ln(4)
        pdf.set_font(base_font, '', 11)
        pdf.cell(0, 8, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True, align='C')
        pdf.ln(10)

        # --- Helper for section titles ---
        def section_title(title: str):
            pdf.set_font(base_font, 'B', 12)
            pdf.cell(0, 8, txt=title.upper(), ln=True, align='L')
            pdf.ln(2)

        # --- Document Summary ---
        summary_text = report_data.get('summary')
        if not summary_text and isinstance(report_data.get('analysis'), dict):
            summary_text = report_data['analysis'].get('raw_text', '')

        if summary_text:
            section_title("Document Summary")
            pdf.set_font(base_font, '', 10)
            pdf.multi_cell(0, 5, clean_text_for_pdf(summary_text))
            pdf.ln(5)

        # --- Risk Assessment ---
        if report_data.get('riskScore'):
            section_title("Risk Assessment")
            risk = report_data['riskScore']
            pdf.set_font(base_font, '', 10)
            pdf.cell(0, 6, f"Score: {risk.get('score', 'N/A')} / {risk.get('max', 'N/A')}", ln=True)
            pdf.cell(0, 6, f"Level: {risk.get('label', 'N/A')}", ln=True)
            pdf.ln(5)

        # --- SWOT Analysis ---
        analysis = report_data.get('analysis', {})
        if isinstance(analysis, dict):
            mapping = {
                "strengths": "Strengths",
                "weaknesses": "Weaknesses",
                "opportunities": "Opportunities",
                "threats": "Threats"
            }

            if not any(k in analysis for k in mapping) and analysis.get('raw_text'):
                section_title("Analysis")
                pdf.set_font(base_font, '', 10)
                pdf.multi_cell(0, 5, clean_text_for_pdf(analysis.get('raw_text', '')))
                pdf.ln(5)
            else:
                for key, label in mapping.items():
                    if analysis.get(key):
                        section_title(label)
                        pdf.set_font(base_font, '', 10)
                        for item in analysis[key]:
                            if isinstance(item, dict) and item.get('text'):
                                pdf.multi_cell(0, 5, f"• {clean_text_for_pdf(item['text'])}")
                        pdf.ln(3)

        # --- Critical Red Flags ---
        if report_data.get('criticalFlags'):
            section_title("Critical Red Flags")
            pdf.set_font(base_font, '', 10)
            for flag in report_data['criticalFlags']:
                if isinstance(flag, dict):
                    title = clean_text_for_pdf(flag.get('title', ''))
                    explanation = clean_text_for_pdf(flag.get('explanation', ''))
                    pdf.set_font(base_font, 'B', 10)
                    pdf.multi_cell(0, 5, f"⚠ {title}")
                    pdf.set_font(base_font, '', 9)
                    pdf.multi_cell(0, 5, explanation)
                    if flag.get('source'):
                        pdf.set_font(base_font, 'I', 9)
                        pdf.cell(0, 5, f"Source: {clean_text_for_pdf(flag['source'])}", ln=True)
                    pdf.ln(4)

        # --- Negotiation Points ---
        if report_data.get('negotiationPoints'):
            section_title("Negotiation Action Plan")
            pdf.set_font(base_font, '', 10)
            for point in report_data['negotiationPoints']:
                if isinstance(point, dict):
                    pdf.set_font(base_font, 'B', 10)
                    pdf.multi_cell(0, 5, f"- {clean_text_for_pdf(point.get('title', ''))}")
                    pdf.set_font(base_font, '', 9)
                    if point.get('risk'):
                        pdf.multi_cell(0, 5, f"Risk: {clean_text_for_pdf(point['risk'])}")
                    if point.get('example'):
                        pdf.multi_cell(0, 5, f"Suggestion: {clean_text_for_pdf(point['example'])}")
                    pdf.ln(4)

        # --- Save ---
        tmp_pdf = tempfile.mktemp(suffix=".pdf")
        pdf.output(tmp_pdf)
        logging.info(f"Saved CovenantAI PDF to {tmp_pdf}")

        # --- Bundle Option (PDF + TXT ZIP) ---
        if bundle:
            raw_text = ''
            if isinstance(report_data.get('analysis'), dict):
                raw_text = report_data['analysis'].get('raw_text', '')
            if not raw_text:
                raw_text = json.dumps(report_data, ensure_ascii=False, indent=2)

            txt_path = tempfile.mktemp(suffix='.txt')
            with open(txt_path, 'w', encoding='utf-8') as tf:
                tf.write(raw_text)

            zip_path = tempfile.mktemp(suffix='.zip')
            import zipfile
            with zipfile.ZipFile(zip_path, 'w', compression=zipfile.ZIP_DEFLATED) as zf:
                zf.write(tmp_pdf, arcname='covenantai-report.pdf')
                zf.write(txt_path, arcname='covenantai-raw.txt')

            return FileResponse(zip_path, media_type='application/zip', filename='covenantai-bundle.zip')

        return FileResponse(
            tmp_pdf,
            media_type='application/pdf',
            filename='covenantai-report.pdf',
            headers={"Content-Disposition": "attachment; filename=covenantai-report.pdf"}
        )

    except Exception as e:
        logging.error(f"CovenantAI PDF export error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to generate PDF: {e}")


# Export document analysis as PDF (Covenant-style: black & white, symmetric, two-column SWOT)
@api_router.get("/documents/{document_id}/export-pdf")
async def export_document_pdf(document_id: str):
    """Export document analysis as a downloadable PDF report (minimal, symmetric, monochrome)."""
    logging.info(f"Starting PDF export for document: {document_id}")

    try:
        import os
        import tempfile
        from datetime import datetime
        import pymongo
        from fastapi import HTTPException
        from fastapi.responses import FileResponse
        from fpdf import FPDF
        import logging

        # --- Fetch document from DB ---
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
        document = sync_db.documents.find_one({"id": document_id})
        sync_client.close()

        if not document:
            logging.error(f"Document not found: {document_id}")
            raise HTTPException(status_code=404, detail="Document not found")

        if document.get('analysis_status') != 'completed':
            logging.error(f"Document analysis not completed: {document_id}, status: {document.get('analysis_status')}")
            raise HTTPException(status_code=400, detail="Document analysis not completed yet")

        # Retrieve analysis sections from the document
        executive_summary = document.get('executive_summary') or ''
        plain_english = document.get('plain_english') or ''
        risk_assessment = document.get('risk_assessment') or ''
        recommendations = document.get('recommendations') or ''
        key_clauses = document.get('key_clauses') or []
        analysis = document.get('analysis') or {}
        analysis_text = document.get('summary') or ''

        logging.info(f"Retrieved document data. Summary length: {len(analysis_text)}")

        # --- Helper: Clean text for PDF (unicode-preserving, safe) ---
        def clean_text_for_pdf(text: str) -> str:
            if not text:
                return ""
            t = str(text)
            # Remove NULs which can break some PDF writers
            t = t.replace('\x00', '')
            # Replace bullet points with asterisks
            t = t.replace('•', '*').replace('\u2022', '*')
            # Replace en/em dashes with hyphens
            t = t.replace('–', '-').replace('—', '-')
            # Normalize smart quotes to ASCII equivalents
            t = t.replace('\u2018', "'").replace('\u2019', "'")
            t = t.replace('\u201C', '"').replace('\u201D', '"')
            # Trim and collapse trailing whitespace
            return t.strip()

        # --- Monochrome PDF class with header/footer & helpers ---
        class MonochromePDF(FPDF):
            def header(self):
                # Title (centered), thin divider line
                self.set_font(self.base_font, 'B', 16)
                self.set_text_color(0, 0, 0)
                self.cell(0, 10, "LEGAL DOCUMENT ANALYSIS REPORT", ln=True, align='C')
                self.ln(2)
                self.set_line_width(0.2)
                self.set_draw_color(0, 0, 0)
                self.line(15, 25, 195, 25)
                self.ln(6)

            def footer(self):
                # Footer with timestamp + page number
                self.set_y(-15)
                self.set_font(self.base_font, '', 8)
                self.set_text_color(0, 0, 0)
                ts = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                self.cell(0, 10, f"Generated {ts}  |  Page {self.page_no()}", align='C')

            def section_title(self, title: str):
                # Left-aligned, uppercased section title with a short divider
                self.set_font(self.base_font, 'B', 12)
                self.cell(0, 8, txt=title, ln=True, align='L')
                self.ln(2)
                self.set_line_width(0.05)
                # short centered divider
                x = self.get_x()
                y = self.get_y()
                self.line(60, y, 150, y)
                self.ln(6)

            def write_paragraph(self, text: str, size: float = 10, leading: float = 5):
                self.set_font(self.base_font, '', size)
                self.set_text_color(0, 0, 0)
                self.multi_cell(0, leading, txt=text)
                self.ln(4)

            def two_column_lists(self, left_items: list, right_items: list, left_title: str = "", right_title: str = ""):
                """Render two lists side-by-side. left_items/right_items are lists of strings."""
                # column widths and start positions
                page_w = self.w - self.l_margin - self.r_margin
                col_w = (page_w - 8) / 2  # 8 units gap between columns
                x_start = self.get_x()
                y_start = self.get_y()

                max_rows = max(len(left_items), len(right_items))

                # Optional small bold titles inside columns
                if left_title:
                    self.set_font(self.base_font, 'B', 11)
                    self.multi_cell(col_w, 5, left_title, border=0)
                if right_title:
                    # move to right column title position
                    self.set_xy(x_start + col_w + 8, y_start)
                    self.set_font(self.base_font, 'B', 11)
                    self.multi_cell(col_w, 5, right_title, border=0)
                self.ln(2)

                # Reset to start of list rows
                y_row = self.get_y()
                for i in range(max_rows):
                    # Left column
                    self.set_xy(x_start, y_row)
                    if i < len(left_items):
                        self.set_font(self.base_font, '', 10)
                        self.multi_cell(col_w, 5, f"• {clean_text_for_pdf(left_items[i])}")
                    else:
                        # empty spacer
                        self.multi_cell(col_w, 5, "")

                    # Right column
                    self.set_xy(x_start + col_w + 8, y_row)
                    if i < len(right_items):
                        self.set_font(self.base_font, '', 10)
                        self.multi_cell(col_w, 5, f"• {clean_text_for_pdf(right_items[i])}")
                    else:
                        self.multi_cell(col_w, 5, "")

                    # Advance y_row to the next line (max of current y positions)
                    y_row = max(self.get_y(), y_row + 5)

                # Move cursor below columns
                self.set_y(y_row + 6)

        # --- Setup PDF object and fonts ---
        pdf = MonochromePDF(format='A4')
        pdf.set_auto_page_break(auto=True, margin=15)
        # Register DejaVu if present; fallback to Arial
        try:
            dejavu_path = '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'
            if os.path.exists(dejavu_path):
                # Add regular and bold variants (both map to same file if no separate bold)
                pdf.add_font('DejaVu', '', dejavu_path, uni=True)
                try:
                    pdf.add_font('DejaVu', 'B', dejavu_path, uni=True)
                except Exception:
                    pass
                base_font = 'DejaVu'
            else:
                base_font = 'Arial'
        except Exception:
            base_font = 'Arial'

        # Attach base_font attribute for class methods to use
        pdf.base_font = base_font

        # Start first page (header uses pdf.base_font)
        pdf.add_page()
        pdf.set_left_margin(15)
        pdf.set_right_margin(15)

        # --- Header meta (centered) ---
        pdf.set_font(base_font, 'B', 14)
        pdf.cell(0, 8, "COVENANTAI REPORT", ln=True, align='C')
        pdf.ln(2)
        pdf.set_font(base_font, '', 10)
        pdf.cell(0, 6, f"Document: {document.get('filename', 'Unknown')}", ln=True, align='C')
        pdf.cell(0, 6, f"Analysis Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True, align='C')
        pdf.ln(8)
        pdf.line(15, pdf.get_y(), 195, pdf.get_y())
        pdf.ln(8)

        # --- Executive / Plain English Summary ---
        if executive_summary or plain_english:
            pdf.section_title("Summary")
            if executive_summary:
                clean_exec = clean_text_for_pdf(executive_summary)
                if len(clean_exec) > 3000:
                    clean_exec = clean_exec[:3000] + " ...[truncated]"
                pdf.write_paragraph(clean_exec, size=10, leading=5)
            if plain_english:
                pdf.write_paragraph(clean_text_for_pdf(plain_english), size=10, leading=5)

        # --- Risk Assessment ---
        if risk_assessment:
            pdf.section_title("Risk Assessment")
            pdf.write_paragraph(clean_text_for_pdf(risk_assessment), size=10, leading=5)

        # --- Recommendations ---
        if recommendations:
            pdf.section_title("Recommendations")
            pdf.write_paragraph(clean_text_for_pdf(recommendations), size=10, leading=5)

        # --- SWOT / Two-column rendering if strengths/weaknesses exist ---
        strengths = []
        weaknesses = []
        if isinstance(analysis, dict):
            # Support both structured list items and simple lists of strings
            if analysis.get('strengths'):
                for it in analysis['strengths']:
                    if isinstance(it, dict):
                        strengths.append(it.get('text', ''))
                    else:
                        strengths.append(str(it))
            if analysis.get('weaknesses'):
                for it in analysis['weaknesses']:
                    if isinstance(it, dict):
                        weaknesses.append(it.get('text', ''))
                    else:
                        weaknesses.append(str(it))

        if strengths or weaknesses:
            pdf.section_title("SWOT (Strengths vs Weaknesses)")
            # Use two-column rendering
            pdf.two_column_lists(left_items=strengths, right_items=weaknesses,
                                 left_title="Strengths", right_title="Weaknesses")

        # --- Key Clauses ---
        if key_clauses:
            pdf.section_title("Key Clauses Analysis")
            for clause in key_clauses:
                if clause.get('clause'):
                    pdf.set_font(base_font, 'B', 10)
                    pdf.multi_cell(0, 5, f"Clause: {clean_text_for_pdf(clause['clause'][:200])}")
                    if clause.get('explanation'):
                        pdf.set_font(base_font, '', 9)
                        pdf.multi_cell(0, 4.5, clean_text_for_pdf(clause.get('explanation', '')))
                    pdf.ln(3)

        # --- Save PDF to temporary file ---
        tmp_file_path = tempfile.mktemp(suffix='.pdf')
        pdf.output(tmp_file_path)
        logging.info(f"PDF saved successfully to: {tmp_file_path}")

        # Return PDF file
        filename = f"covenant_report_{document_id}.pdf"
        logging.info(f"Returning PDF file: {filename}")
        return FileResponse(
            tmp_file_path,
            media_type='application/pdf',
            filename=filename,
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )

    except HTTPException:
        # re-raise known HTTP exceptions
        raise
    except Exception as e:
        logging.error(f"PDF export error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to generate PDF report: {str(e)}")



# Then include the router in the main app
app.include_router(api_router)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    # Close any global MongoDB clients if present
    try:
        if MongoDBState.sync_client:
            MongoDBState.sync_client.close()
    except Exception:
        pass
    try:
        if MongoDBState.async_client:
            MongoDBState.async_client.close()
    except Exception:
        pass

if __name__ == "__main__":
    import uvicorn, os
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)

